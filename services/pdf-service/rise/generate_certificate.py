from docx import Document
import fitz  # PyMuPDF
from typing import Dict
from .font_utils import calculate_optimal_font_size_with_line_breaks
import unidecode
import ftfy
import chardet
#CraftApp - Copy
def parse_excel_adjustment(value):
    """Parse Excel adjustment value (e.g., '-1', '+2', '3', '', None) and return numeric value."""
    # Handle None, empty string, or whitespace-only values
    if value is None or (isinstance(value, str) and value.strip() == ''):
        return 0
    
    value = str(value).strip()
    
    # Handle negative values
    if value.startswith('-'):
        try:
            return -float(value[1:])
        except (ValueError, TypeError):
            return 0
    
    # Handle positive values (with or without +)
    if value.startswith('+'):
        value = value[1:]
    
    try:
        return float(value)
    except (ValueError, TypeError):
        return 0

def parse_excel_font_size(value):
    """Parse Excel font size value (e.g., '-1', '+2', '3', '', None) and return numeric value."""
    # Handle None, empty string, or whitespace-only values
    if value is None or (isinstance(value, str) and value.strip() == ''):
        return 0
    
    value = str(value).strip()
    
    # Handle negative values
    if value.startswith('-'):
        try:
            return -float(value[1:])
        except (ValueError, TypeError):
            return 0
    
    # Handle positive values (with or without +)
    if value.startswith('+'):
        value = value[1:]
    
    try:
        return float(value)
    except (ValueError, TypeError):
        return 0

def safe_insert_text(page, position, text, **kwargs):
    """Safely insert text handling Unicode characters that might cause ByteString errors."""
    try:
        # 🔍 DEBUG: Analyze text at start
        # Check for any apostrophe-like character (regular apostrophe or smart quotes)
        has_regular_apostrophe = "'" in text  # U+0027
        # Check for smart quotes using their Unicode code points
        has_right_single_quote = '\u2019' in text  # RIGHT SINGLE QUOTATION MARK
        has_left_single_quote = '\u2018' in text   # LEFT SINGLE QUOTATION MARK
        has_apostrophe = has_regular_apostrophe or has_right_single_quote or has_left_single_quote
        
        print(f"🔍 [SAFE_INSERT] Analyzing text: '{text[:50]}...'")
        print(f"🔍 [SAFE_INSERT] Regular apostrophe ('): {has_regular_apostrophe}")
        print(f"🔍 [SAFE_INSERT] Right quote (U+2019): {has_right_single_quote}")
        print(f"🔍 [SAFE_INSERT] Left quote (U+2018): {has_left_single_quote}")
        print(f"🔍 [SAFE_INSERT] Has apostrophe-like character: {has_apostrophe}")
        
        # First try with the original text
        # 🔍 DEBUG: Check for special dash characters
        if '–' in text or '—' in text or '·' in text:
            for i, char in enumerate(text):
                if char in ['–', '—', '·', '-', '−']:
                    print(f"🔍 [SAFE_INSERT] Found dash/dot at pos {i}: '{char}' (Unicode: {ord(char)}) in text: '{text[:50]}...'")
        
        # 🔍 DEBUG: Check for apostrophes (regular or smart quotes)
        if has_apostrophe:
            print(f"🔍 [SAFE_INSERT] ✓ Apostrophe-like character found in text!")
            for i, char in enumerate(text):
                if char in ["'", '\u2019', '\u2018']:
                    print(f"🔍 [SAFE_INSERT] Found apostrophe at pos {i}: '{char}' (Unicode: {ord(char)}) in text: '{text[:50]}...'")
        else:
            print(f"🔍 [SAFE_INSERT] ✗ No apostrophe found in text")
        
        # ✅ ENHANCED: Detect special characters that need special handling
        special_chars = ['–', '—', '·', '\u2019', '\u2018', '"', '"', '…', '€', '£', '¥', '©', '®', '™']
        has_special_chars = any(char in text for char in special_chars)
        
        if has_special_chars:
            print(f"🔍 [SAFE_INSERT] Special characters detected in text: '{text[:50]}...'")
            
            # Fix text encoding issues first
            fixed_text = ftfy.fix_text(text)
            if fixed_text != text:
                print(f"🔍 [SAFE_INSERT] Text encoding fixed: '{text[:30]}...' → '{fixed_text[:30]}...'")
                text = fixed_text
            
            # ✅ CHANGED: Keep original font (Times-Bold) - special chars will be replaced with ASCII equivalents below
            # No font switching needed since we replace special characters with ASCII equivalents
            print(f"🔍 [SAFE_INSERT] Keeping original font '{kwargs.get('fontname', 'Times-Roman')}' - special chars will be replaced with ASCII equivalents")
        elif has_apostrophe:
            # ✅ CHANGED: Keep original font (Times-Bold) - apostrophes will be replaced with ASCII equivalents below
            # No font switching needed since we replace apostrophes with ASCII equivalents
            print(f"🔍 [SAFE_INSERT] Keeping original font '{kwargs.get('fontname', 'Times-Roman')}' - apostrophes will be replaced with ASCII equivalents")
        
        # Replace problematic Unicode characters with ASCII equivalents for Times font compatibility
        text = text.replace('–', '-')  # En dash (U+2013) → hyphen
        text = text.replace('—', '-')  # Em dash (U+2015) → hyphen
        
        # Replace smart quotes with straight apostrophes for better font compatibility
        text = text.replace('\u2019', "'")  # Right single quotation mark (U+2019) → apostrophe
        text = text.replace('\u2018', "'")  # Left single quotation mark (U+2018) → apostrophe
        text = text.replace('\u201D', '"')  # Right double quotation mark (U+201D) → straight quote
        text = text.replace('\u201C', '"')  # Left double quotation mark (U+201C) → straight quote
        
        page.insert_text(position, text, **kwargs)
    except Exception as e:
        if "ByteString" in str(e) or "character at index" in str(e):
            print(f"⚠️ [CERTIFICATE] Unicode text error, using safe encoding: {e}")
            # Try with UTF-8 encoding that ignores problematic characters
            safe_text = text.encode('utf-8', errors='ignore').decode('utf-8')
            try:
                page.insert_text(position, safe_text, **kwargs)
            except Exception as e2:
                print(f"⚠️ [CERTIFICATE] UTF-8 encoding failed, using ASCII fallback: {e2}")
                # Final fallback: Use unidecode for ASCII conversion
                ascii_text = unidecode.unidecode(text)
                print(f"🔍 [SAFE_INSERT] ASCII fallback text: '{ascii_text}'")
                page.insert_text(position, ascii_text, **kwargs)
        else:
            # Re-raise if it's not a Unicode/ByteString error
            raise e

# ISO Standards Mapping - Convert short names to full versions with years
ISO_STANDARDS_MAPPING = {
    # Quality & Management
    "ISO 9001": "ISO 9001:2015",
    "9001": "ISO 9001:2015",
    "ISO 14001": "ISO 14001:2015", 
    "14001": "ISO 14001:2015",
    "ISO 45001": "ISO 45001:2018",
    "45001": "ISO 45001:2018",
    "ISO 50001": "ISO 50001:2018",
    "50001": "ISO 50001:2018",
    "ISO 31000": "ISO 31000:2018",
    "31000": "ISO 31000:2018",
    
    # Food Safety
    "ISO 22000": "ISO 22000:2018",
    "22000": "ISO 22000:2018",
    "ISO/TS 22002-1": "ISO/TS 22002-1:2009",
    "22002-1": "ISO/TS 22002-1:2009",
    "ISO 22005": "ISO 22005:2007",
    "22005": "ISO 22005:2007",
    
    # Laboratory & Testing
    "ISO/IEC 17025": "ISO/IEC 17025:2017",
    "17025": "ISO/IEC 17025:2017",
    "ISO 15189": "ISO 15189:2022",
    "15189": "ISO 15189:2022",
    
    # Information Security & IT
    "ISO/IEC 27001": "ISO/IEC 27001:2022",
    "27001": "ISO/IEC 27001:2022",
    "ISO/IEC 27002": "ISO/IEC 27002:2022",
    "27002": "ISO/IEC 27002:2022",
    "ISO/IEC 20000-1": "ISO/IEC 20000-1:2018",
    "20000-1": "ISO/IEC 20000-1:2018",
    "ISO 22301": "ISO 22301:2019",
    "22301": "ISO 22301:2019",
    
    # Manufacturing & Industrial
    "ISO 13485": "ISO 13485:2016",
    "13485": "ISO 13485:2016",
    "IATF 16949": "IATF 16949:2016",
    "16949": "IATF 16949:2016",
    "ISO 3834-2": "ISO 3834-2:2021",
    "3834-2": "ISO 3834-2:2021",
    
    # Environment & Sustainability
    "ISO 14064-1": "ISO 14064-1:2018",
    "14064-1": "ISO 14064-1:2018",
    "ISO 14046": "ISO 14046:2014",
    "14046": "ISO 14046:2014",
    "ISO 20121": "ISO 20121:2012",
    "20121": "ISO 20121:2012",
    
    # Asset, Facility, and Supply Chain
    "ISO 55001": "ISO 55001:2014",
    "55001": "ISO 55001:2014",
    "ISO 28000": "ISO 28000:2022",
    "28000": "ISO 28000:2022",
    
    # Aerospace
    "AS 9100D": "AS 9100D:2016",
    "9100D": "AS 9100D:2016",
    
    # Other Notable Standards
    "ISO 37001": "ISO 37001:2016",
    "37001": "ISO 37001:2016",
    "ISO 19600": "ISO 19600:2014",
    "19600": "ISO 19600:2014",
    "ISO 29993": "ISO 29993:2017",
    "29993": "ISO 29993:2017",
}

# ISO Standards Code Mapping - For certification codes
ISO_STANDARDS_CODES = {
    "ISO 9001:2015": "CM-MS-7842",
    "ISO 14001:2015": "CM-MS-7836", 
    "ISO 45001:2018": "CM-MS-7832",
    "ISO 22000:2018": "CM-MS-7822",
    "ISO/IEC 27001:2022": "CM-MS-7820",
    "ISO 37001:2016": "CM-MS-7804",
    "ISO 22301:2019": "CM-MS-7807",
    "ISO 50001:2018": "CM-MS-7814",
    "ISO 20001:2018": "CM-MS-7811",
}

# ISO Standards Descriptions Mapping - For separate use
ISO_STANDARDS_DESCRIPTIONS = {
    "ISO 9001:2015": "Quality Management System",
    "ISO 14001:2015": "Environmental Management System",
    "ISO 45001:2018": "Occupational Health & Safety Management System",
    "ISO 50001:2018": "Energy Management System",
    "ISO 31000:2018": "Risk Management Guidelines",
    "ISO 22000:2018": "Food Safety Management System",
    "ISO/TS 22002-1:2009": "Prerequisite programs on food safety",
    "ISO 22005:2007": "Traceability in the feed and food chain",
    "ISO/IEC 17025:2017": "Testing and Calibration Laboratories",
    "ISO 15189:2022": "Medical Laboratories – Quality and Competence",
    "ISO/IEC 27001:2022": "Information Security Management System",
    "ISO/IEC 27002:2022": "Information Security Controls",
    "ISO/IEC 20000-1:2018": "IT Service Management System",
    "ISO 22301:2019": "Business Continuity Management System",
    "ISO 13485:2016": "Medical Devices – Quality Management System",
    "IATF 16949:2016": "Automotive Quality Management System",
    "ISO 3834-2:2021": "Quality requirements for fusion welding",
    "ISO 14064-1:2018": "Greenhouse Gases",
    "ISO 14046:2014": "Water Footprint",
    "ISO 20121:2012": "Event Sustainability Management System",
    "ISO 55001:2014": "Asset Management System",
    "ISO 28000:2022": "Security Management Systems for Supply Chain",
    "AS 9100D:2016": "Aerospace Quality (based on ISO 9001:2015)",
    "ISO 37001:2016": "Anti-bribery Management System",
    "ISO 19600:2014": "Compliance Management System",
    "ISO 29993:2017": "Learning Services",
}

# Spanish ISO Standards Descriptions Mapping
ISO_STANDARDS_DESCRIPTIONS_SPANISH = {
    "ISO 9001:2015": "el Sistema de Gestión de Calidad",
    "ISO 14001:2015": "el Sistema de gestión ambiental",
    "ISO 45001:2018": "el Sistema de gestión de seguridad y salud en el trabajo",
    "ISO 50001:2018": "el Sistema de Gestión Energética",
    "ISO 31000:2018": "Directrices para la Gestión del Riesgo",
    "ISO 22000:2018": "el Sistema de Gestión de Seguridad Alimentaria",
    "ISO/TS 22002-1:2009": "Programas de prerrequisitos en seguridad alimentaria",
    "ISO 22005:2007": "Trazabilidad en la cadena alimentaria y de piensos",
    "ISO/IEC 17025:2017": "Laboratorios de Ensayo y Calibración",
    "ISO 15189:2022": "Laboratorios Médicos – Calidad y Competencia",
    "ISO/IEC 27001:2022": "Sistema de Gestión de Seguridad de la Información",
    "ISO/IEC 27002:2022": "Controles de Seguridad de la Información",
    "ISO/IEC 20000-1:2018": "Sistema de Gestión de Servicios de TI",
    "ISO 22301:2019": "Sistema de Gestión de Continuidad del Negocio",
    "ISO 13485:2016": "Dispositivos Médicos – Sistema de Gestión de Calidad",
    "IATF 16949:2016": "Sistema de Calidad Automotriz",
    "ISO 3834-2:2021": "Requisitos de calidad para soldadura por fusión",
    "ISO 14064-1:2018": "Gases de Efecto Invernadero",
    "ISO 14046:2014": "Huella Hídrica",
    "ISO 20121:2012": "Sistema de Gestión de Sostenibilidad de Eventos",
    "ISO 55001:2014": "Sistema de Gestión de Activos",
    "ISO 28000:2022": "Sistemas de Gestión de Seguridad para la Cadena de Suministro",
    "AS 9100D:2016": "Calidad Aeroespacial (basado en ISO 9001:2015)",
    "ISO 37001:2016": "el Sistema de gestión antisoborno",
    "ISO 19600:2014": "Sistema de Gestión de Cumplimiento",
    "ISO 29993:2017": "Servicios de Aprendizaje",
}

def expand_iso_standard(iso_text: str) -> str:
    """Expand ISO standard name to full version with year if available."""
    if not iso_text:
        return iso_text
    
    # Clean the input text
    cleaned_text = iso_text.strip()
    
    # First, try exact match
    if cleaned_text in ISO_STANDARDS_MAPPING:
        return ISO_STANDARDS_MAPPING[cleaned_text]
    
    # If no exact match, try to find partial matches
    # This handles cases where users might enter variations
    for short_name, full_name in ISO_STANDARDS_MAPPING.items():
        # Check if the input contains the standard number
        if short_name.lower() in cleaned_text.lower():
            return full_name
    
    # If still no match, try to extract just the number and match
    # This handles cases like "37001" when we have "37001" in mapping
    import re
    number_match = re.search(r'(\d+(?:-\d+)?)', cleaned_text)
    if number_match:
        number = number_match.group(1)
        if number in ISO_STANDARDS_MAPPING:
            return ISO_STANDARDS_MAPPING[number]
    
    # If no match found, return original text
    return iso_text

def get_iso_standard_code(iso_standard: str) -> str:
    """
    Get the certification code for a given ISO standard.
    Example: "ISO 9001:2015" -> "CM-MS-7842"
    Returns empty string if no code mapping exists.
    """
    if not iso_standard:
        return ""
    
    # First expand the ISO standard if it's a short name
    expanded_iso = expand_iso_standard(iso_standard)
    
    # Look up the code in the mapping
    return ISO_STANDARDS_CODES.get(expanded_iso, "")

def parse_word_form(docx_path: str) -> Dict[str, str]:
    """Parse the first table in a Word document and extract required fields."""
    
    doc = Document(docx_path)
    
    if len(doc.tables) == 0:
        raise Exception("No tables found in document")
    
    table = doc.tables[0]
    
    data = {}
    for i, row in enumerate(table.rows):
        if len(row.cells) == 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            data[key] = value
        else:
            continue
    
    # Get ISO Standard and expand it to full version with year
    iso_standard = data.get("ISO Standard Required", "").splitlines()[-1]
    expanded_iso = expand_iso_standard(iso_standard)
    
    result = {
        "Company Name": data.get("Company Name", ""),
        "Address": data.get("Address", ""),
        "ISO Standard": expanded_iso,
        "Scope": data.get("Scope", "")
    }
    
    return result

def preprocess_image_for_ocr(image):
    """Preprocess image to improve OCR accuracy."""
    import cv2
    import numpy as np
    
    # Convert to grayscale
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Apply noise reduction
    denoised = cv2.medianBlur(gray, 3)
    
    # Apply thresholding
    _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Apply morphological operations
    kernel = np.ones((1, 1), np.uint8)
    processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
    
    return processed

def parse_ocr_text_as_table(text):
    """Parse OCR text to detect table structure."""
    lines = text.split('\n')
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Look for field:value patterns
        if ':' in line:
            parts = line.split(':', 1)
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()
                table_data.append([key, value])
        else:
            # This might be a continuation line
            table_data.append(['', line])
    
    return table_data

def extract_fields_from_ocr_text(text):
    """Extract fields from OCR text using pattern matching."""
    import re
    
    data = {}
    
    # Define patterns for field extraction
    patterns = {
        "Company Name": [
            r"Company Name[:\s]+(.+?)(?=\n\s*(?:Address|ISO Standard Required|Scope)|$)",
            r"Company[:\s]+(.+?)(?=\n\s*(?:Address|ISO Standard Required|Scope)|$)"
        ],
        "Address": [
            r"Address[:\s]+(.+?)(?=\n\s*(?:ISO Standard Required|Scope)|$)",
            r"Location[:\s]+(.+?)(?=\n\s*(?:ISO Standard Required|Scope)|$)"
        ],
        "ISO Standard Required": [
            r"ISO Standard Required[:\s]+(.+?)(?=\n\s*(?:Scope)|$)",
            r"ISO Standard[:\s]+(.+?)(?=\n\s*(?:Scope)|$)"
        ],
        "Scope": [
            r"Scope[:\s]+(.+?)(?=\n\s*(?:Company Name|Address|ISO Standard Required)|$)"
        ]
    }
    
    for field_name, field_patterns in patterns.items():
        for pattern in field_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                value = match.group(1).strip()
                if value:
                    data[field_name] = value
                    break
    
    return data

def extract_from_images(image_path: str) -> Dict[str, str]:
    """Extract fields from image using OCR with table detection."""
    import pytesseract
    import cv2
    import numpy as np
    from PIL import Image
    
    print(f"🔍 [IMAGE-DEBUG] Starting OCR extraction from: {image_path}")
    
    try:
        # Load and preprocess image
        image = cv2.imread(image_path)
        if image is None:
            raise Exception(f"Could not load image: {image_path}")
            
        processed_image = preprocess_image_for_ocr(image)
        
        # OCR with table detection
        custom_config = r'--oem 3 --psm 6'  # Table detection mode
        text = pytesseract.image_to_string(processed_image, config=custom_config)
        
        print(f"🔍 [IMAGE-DEBUG] OCR extracted text:")
        print(f"🔍 [IMAGE-DEBUG] {text[:1000]}{'...' if len(text) > 1000 else ''}")
        
        # Try to detect table structure
        table_data = parse_ocr_text_as_table(text)
        if table_data:
            print(f"🔍 [IMAGE-DEBUG] Table structure detected: {len(table_data)} rows")
            return process_table_data(table_data)
        
        # Fallback to text pattern matching
        print(f"🔍 [IMAGE-DEBUG] No table structure found, using pattern matching")
        return extract_fields_from_ocr_text(text)
        
    except Exception as e:
        print(f"🔍 [IMAGE-DEBUG] Error in OCR extraction: {e}")
        raise Exception(f"Failed to extract text from image: {str(e)}")

def process_table_data(table_data):
    """Process table data to extract fields, similar to PDF table processing."""
    data = {}
    last_recognized_field = None
    
    for i, row in enumerate(table_data):
        if len(row) >= 2:
            key = str(row[0]).strip() if row[0] else ""
            value = str(row[1]).strip() if row[1] else ""
            
            print(f"🔍 [IMAGE-DEBUG] Row {i+1}: '{key}' -> '{value[:50]}{'...' if len(value) > 50 else ''}'")
            
            # Check if this is a recognized field
            if key in ['Company Name', 'Address', 'ISO Standard Required', 'Scope']:
                data[key] = value
                last_recognized_field = key
                print(f"🔍 [IMAGE-DEBUG] ✅ Found recognized field '{key}': '{value[:100]}{'...' if len(value) > 100 else ''}'")
            elif key == "" and value and last_recognized_field:
                # This is a continuation line (empty key, has value)
                data[last_recognized_field] += " " + value
                print(f"🔍 [IMAGE-DEBUG] 🔗 Appended continuation to '{last_recognized_field}': '{value[:50]}{'...' if len(value) > 50 else ''}'")
            else:
                print(f"🔍 [IMAGE-DEBUG] ⏭️ Skipping unrecognized field '{key}'")
    
    return data

def parse_pdf_form(pdf_path: str) -> Dict[str, str]:
    """Parse a PDF document or image using hybrid table + text + OCR extraction approach."""
    
    try:
        # Determine file type
        file_extension = pdf_path.lower().split('.')[-1]
        
        if file_extension in ['png', 'jpg', 'jpeg']:
            # Phase 1: Try image extraction with OCR
            print(f"🔍 [PDF-DEBUG] Detected image file, using OCR extraction")
            data = extract_from_images(pdf_path)
        else:
            # Phase 2: Try table extraction first
            data = extract_from_tables(pdf_path)
        
        # Phase 2: If table extraction fails or is incomplete, use text extraction
        if not data or len(data) < 4:
            print(f"🔍 [PDF-DEBUG] Table extraction incomplete ({len(data) if data else 0}/4 fields), trying text extraction...")
            data = extract_from_text(pdf_path)
        
        # If still no data found, raise exception
        if not data:
            raise Exception("No recognizable field patterns found in PDF")
        
        # Get ISO Standard and expand it to full version with year
        iso_standard = data.get("ISO Standard Required", "").splitlines()[-1] if data.get("ISO Standard Required") else ""
        expanded_iso = expand_iso_standard(iso_standard)
        
        result = {
            "Company Name": data.get("Company Name", ""),
            "Address": data.get("Address", ""),
            "ISO Standard": expanded_iso,
            "Scope": data.get("Scope", "")
        }
        
        print(f"🔍 [PDF-DEBUG] Final extracted fields:")
        for key, value in result.items():
            print(f"🔍 [PDF-DEBUG] {key}: '{value[:100]}{'...' if len(value) > 100 else ''}'")
        
        return result
        
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

def extract_from_tables(pdf_path: str) -> Dict[str, str]:
    """Extract fields from PDF using table extraction - fixes contamination issue."""
    
    doc = fitz.open(pdf_path)
    data = {}
    
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Try to extract tables from the page
            tables = page.find_tables()
            
            # Convert TableFinder to list
            tables_list = list(tables)
            
            if tables_list:
                print(f"🔍 [PDF-DEBUG] Found {len(tables_list)} table(s) on page {page_num + 1}")
                
                # Use the first table found
                table = tables_list[0]
                table_data = table.extract()
                
                print(f"🔍 [PDF-DEBUG] Table data extracted: {len(table_data)} rows")
                
                # Process each row as key-value pairs
                last_recognized_field = None  # Track the last recognized field
                
                for i, row in enumerate(table_data):
                    if len(row) >= 2:
                        key = str(row[0]).strip() if row[0] else ""
                        value = str(row[1]).strip() if row[1] else ""
                        
                        print(f"🔍 [PDF-DEBUG] Row {i+1}: '{key}' -> '{value[:50]}{'...' if len(value) > 50 else ''}'")
                        
                        # Check if this is a recognized field
                        if key in ['Company Name', 'Address', 'ISO Standard Required', 'Scope']:
                            data[key] = value
                            last_recognized_field = key  # Track the last recognized field
                            print(f"🔍 [PDF-DEBUG] ✅ Found recognized field '{key}': '{value[:100]}{'...' if len(value) > 100 else ''}'")
                        elif key == "" and value and last_recognized_field:
                            # This is a continuation line (empty key, has value)
                            # Append to the last recognized field
                            data[last_recognized_field] += " " + value
                            print(f"🔍 [PDF-DEBUG] 🔗 Appended continuation to '{last_recognized_field}': '{value[:50]}{'...' if len(value) > 50 else ''}'")
                        else:
                            print(f"🔍 [PDF-DEBUG] ⏭️ Skipping unrecognized field '{key}'")
                
                # If we found data in tables, use it
                if data:
                    print(f"🔍 [PDF-DEBUG] Table extraction successful: {len(data)} fields found")
                    break
            else:
                print(f"🔍 [PDF-DEBUG] No tables found on page {page_num + 1}")
    
    finally:
        doc.close()
    
    return data

def extract_from_text(pdf_path: str) -> Dict[str, str]:
    """Extract fields from PDF using text extraction as fallback."""
    
    doc = fitz.open(pdf_path)
    text = ""
    
    try:
        # Extract all text from PDF
        for page in doc:
            text += page.get_text() + "\n"
        
        print(f"🔍 [PDF-DEBUG] Extracted text from PDF:")
        print(f"🔍 [PDF-DEBUG] {text[:1000]}{'...' if len(text) > 1000 else ''}")
        print(f"🔍 [PDF-DEBUG] ===== END EXTRACTED TEXT =====")
        
        # Simple field extraction without regex - handle multi-line content properly
        lines = text.split('\n')
        data = {}
        current_field = None
        current_value = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this line starts a new field (look for field names followed by colon or at start of line)
            field_names = ['Company Name', 'Address', 'ISO Standard Required', 'Scope']
            is_field_start = False
            detected_field = None
            
            for field_name in field_names:
                if line.startswith(field_name + ':') or line == field_name:
                    is_field_start = True
                    detected_field = field_name
                    break
            
            if is_field_start:
                # Save previous field
                if current_field and current_value:
                    data[current_field] = '\n'.join(current_value).strip()
                    print(f"🔍 [PDF-DEBUG] Saved field '{current_field}': '{data[current_field][:100]}{'...' if len(data[current_field]) > 100 else ''}'")
                
                # Start new field
                current_field = detected_field
                current_value = []
                
                # Extract value from the same line if it exists after the colon
                if ':' in line:
                    value_part = line.split(':', 1)[1].strip()
                    if value_part:
                        current_value.append(value_part)
            elif current_field:
                # This is a continuation of the current field
                current_value.append(line)
        
        # Save last field
        if current_field and current_value:
            data[current_field] = '\n'.join(current_value).strip()
            print(f"🔍 [PDF-DEBUG] Saved final field '{current_field}': '{data[current_field][:100]}{'...' if len(data[current_field]) > 100 else ''}'")
    
    finally:
        doc.close()
    
    return data

def extract_fields_from_text(text: str) -> Dict[str, str]:
    """Extract fields from text using pattern matching as fallback."""
    
    data = {}
    
    # Debug: Print the extracted text to understand the structure
    print(f"🔍 [PDF-DEBUG] Extracted text from PDF:")
    print(f"🔍 [PDF-DEBUG] {text[:500]}{'...' if len(text) > 500 else ''}")
    print(f"🔍 [PDF-DEBUG] ===== END EXTRACTED TEXT =====")
    
    # Define patterns for field extraction - capture until next field or end
    patterns = {
        "Company Name": [
            r"Company Name[:\s]+(.+?)(?=\n\s*(?:Address|ISO Standard Required|Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"Company[:\s]+(.+?)(?=\n\s*(?:Address|ISO Standard Required|Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"Organization[:\s]+(.+?)(?=\n\s*(?:Address|ISO Standard Required|Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)"
        ],
        "Address": [
            r"Address[:\s]+(.+?)(?=\n\s*(?:ISO Standard Required|Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"Location[:\s]+(.+?)(?=\n\s*(?:ISO Standard Required|Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)"
        ],
        "ISO Standard Required": [
            r"ISO Standard Required[:\s]+(.+?)(?=\n\s*(?:Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"ISO Standard[:\s]+(.+?)(?=\n\s*(?:Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"Standard[:\s]+(.+?)(?=\n\s*(?:Scope|Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)"
        ],
        "Scope": [
            r"Scope[:\s]+(.+?)(?=\n\s*(?:Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)",
            r"Scope of Work[:\s]+(.+?)(?=\n\s*(?:Certificate Number|Original Issue Date|Issue Date|Surveillance/Expiry Date|Recertification Date|Initial Registration Date|Surveillance Due Date|Expiry Date|Extra Line)|$)"
        ]
    }
    
    import re
    
    for field_name, field_patterns in patterns.items():
        for pattern in field_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                value = match.group(1).strip()
                if value:
                    data[field_name] = value
                    print(f"🔍 [PDF-DEBUG] Found {field_name}: '{value[:100]}{'...' if len(value) > 100 else ''}'")
                    break
    
    return data

def get_text_height(text: str, fontsize: float, fontname: str, max_width: float) -> float:
    """Estimate the height of a text block when wrapped to fit max_width."""
    font = fitz.Font(fontname=fontname)
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        if font.text_length(test_line, fontsize) <= max_width:
            current_line = test_line
        else:
            lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return len(lines) * fontsize * 1.2  # Approximate line height with spacing

def insert_centered_textbox(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    fontname: str,
    fontsize: float,
    color: tuple
) -> None:
    """Insert text centered both vertically and horizontally in the given rectangle."""
    # Calculate text height and center it vertically
    text_height = get_text_height(text, fontsize, fontname, rect.width)
    start_y = rect.y0 + (rect.height - text_height) / 2
    box = fitz.Rect(rect.x0, start_y, rect.x1, start_y + text_height)
    
    page.insert_textbox(
        box,
        text,
        fontsize=fontsize,
        fontname=fontname,
        color=color,
        align=1  # Centered
    )

def generate_certificate(base_pdf_path: str, output_pdf_path: str, values: Dict[str, str], template_type: str = "standard") -> Dict[str, any]:
    """Generate a certificate PDF by overlaying extracted values onto a template.
    
    Returns:
        Dict containing success status and overflow warnings
    """
    print("CERT DEBUG BUILD: 2025-09-10-14:20")

    
    # Initialize tracking for overflow warnings
    overflow_warnings = []
    doc = fitz.open(base_pdf_path)
    page = doc[0]

    # Company/context snapshot for debugging runs (helps identify Kotec, etc.)
    try:
        company_dbg = (values.get("Company Name") or values.get("company_name") or "").strip()
        country_dbg = (values.get("Country") or values.get("country") or "").strip()
        accred_dbg = (values.get("Accreditation") or values.get("accreditation") or "").strip()
        print(f"[CHECK] Company='{company_dbg}' | Country='{country_dbg}' | Accreditation='{accred_dbg}' | Template='{template_type}'")
    except Exception:
        pass

    # --- Configuration ---
    color = (0, 0, 0)  # Black text
    fontname = "Times-Bold"  # Use Times New Roman Bold font
    
    # ✅ ADDED: Font weight preservation system
    def detect_font_weight(text):
        """
        Detect font weight from text formatting.
        Excel doesn't preserve bold formatting, but we can implement
        a marker system for future enhancement.
        """
        # For now, return default font
        # Future enhancement: Parse Excel formatting or use markers like **bold** or __bold__
        return "Times-Bold"
    
    def process_bold_text(text):
        """
        Process text with bold markers and return segments with font information.
        Returns list of tuples: (text_segment, font_name, is_bold)
        """
        if not text:
            return [(text, "Times-Bold", False)]
        
        segments = []
        current_text = text
        
        # Process **bold** markers
        while '**' in current_text:
            parts = current_text.split('**', 2)
            if len(parts) >= 3:
                # Add normal text before bold
                if parts[0]:
                    segments.append((parts[0], "Times-Roman", False))
                # Add bold text
                segments.append((parts[1], "Times-Bold", True))
                # Continue with remaining text
                current_text = parts[2]
            else:
                break
        
        # Process __bold__ markers
        while '__' in current_text:
            parts = current_text.split('__', 2)
            if len(parts) >= 3:
                # Add normal text before bold
                if parts[0]:
                    segments.append((parts[0], "Times-Roman", False))
                # Add bold text
                segments.append((parts[1], "Times-Bold", True))
                # Continue with remaining text
                current_text = parts[2]
            else:
                break
        
        # Add any remaining normal text
        if current_text:
            segments.append((current_text, "Times-Roman", False))
        
        # If no bold markers found, return original text with default font
        if not segments:
            segments.append((text, "Times-Roman", False))
        
        return segments

    def get_font_for_text(text, default_font="Times-Bold"):
        """
        Get appropriate font for text based on content analysis.
        This is now a legacy function - use process_bold_text for full processing.
        """
        # Check for bold markers
        if '**' in text or '__' in text:
            return "Times-Bold"  # Will be processed by process_bold_text
        else:
            return default_font

    def render_mixed_format_text(page, position, text, font_size, color, max_width=None):
        """
        Render text with mixed bold/normal formatting at the specified position.
        Returns the total width used for positioning calculations.
        """
        if not text:
            return 0
        
        segments = process_bold_text(text)
        current_x = position[0]
        total_width = 0
        
        for segment_text, font_name, is_bold in segments:
            if not segment_text:
                continue
                
            # Calculate text width
            font_obj = fitz.Font(fontname=font_name)
            text_width = font_obj.text_length(segment_text, font_size)
            
            # Check if we need to wrap (if max_width is specified)
            if max_width and current_x + text_width > position[0] + max_width:
                # For now, just render what fits - could implement word wrapping here
                pass
            
            # Render the text segment
            safe_insert_text(
                page,
                (current_x, position[1]),
                segment_text,
                fontsize=font_size,
                fontname=font_name,
                color=color
            )
            
            # Move position for next segment
            current_x += text_width
            total_width += text_width
        
        return total_width

    # ✅ ADDED: Logo processing
    logo_lookup = values.get("logo_lookup", {})
    logo_filename = values.get("Logo", "").strip()
    
    # Process logo if specified and available
    logo_image = None
    if logo_lookup and len(logo_lookup) > 0:
        try:
            # Convert uploaded file to PIL Image
            from PIL import Image
            import io
            
            # Use exact match if available, otherwise try partial match, otherwise use first available
            logo_file = None
            if logo_filename and logo_filename in logo_lookup:
                logo_file = logo_lookup[logo_filename]
                print(f"🔍 [CERTIFICATE] Using exact logo match: {logo_filename}")
            elif logo_filename:
                # Try partial match (filename without extension)
                for filename, file in logo_lookup.items():
                    if logo_filename in filename or filename.split('.')[0] == logo_filename:
                        logo_file = file
                        print(f"🔍 [CERTIFICATE] Using partial logo match: '{logo_filename}' → '{filename}'")
                        break
            
            if not logo_file:
                # Use first available logo file, but check if it's a valid image format
                for filename, file in logo_lookup.items():
                    # Check if file has valid image extension
                    if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                        logo_file = file
                        print(f"🔍 [CERTIFICATE] Using first valid image file: {filename}")
                        break
                
                if not logo_file:
                    print(f"⚠️ [CERTIFICATE] No valid image files found in logo_lookup: {list(logo_lookup.keys())}")
            
            if logo_file and hasattr(logo_file, 'file'):
                # Reset file pointer
                logo_file.file.seek(0)
                # Read file content
                logo_content = logo_file.file.read()
                # Convert to PIL Image
                logo_image = Image.open(io.BytesIO(logo_content))
                print(f"✅ [CERTIFICATE] Logo image loaded successfully")
            else:
                logo_image = None
                print(f"⚠️ [CERTIFICATE] Logo file has no file attribute or is None")
        except Exception as logo_error:
            logo_image = None
            print(f"❌ [CERTIFICATE] Error loading logo image: {logo_error}")
    else:
        logo_image = None
        print(f"🔍 [CERTIFICATE] No logo to process: logo_filename='{logo_filename}', logo_lookup_count={len(logo_lookup) if logo_lookup else 0}")
    
    # Standard template coordinates (current)
    standard_coords = {
        "management_system": fitz.Rect(87.9, 185, 580, 226.6),
        "Company Name and Address": fitz.Rect(87.9, 239, 580, 315),  # Fixed: Match soft copy coordinates
        "ISO Standard": fitz.Rect(194.9, 334, 460.3, 370),
        "Scope": {
            "short": fitz.Rect(87.9, 386, 580, 475),    # <24 lines
            "long": fitz.Rect(87.9, 373, 580, 486)      # 24-30 lines
        },
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # Large template coordinates (for >11 lines)
    # All generation types use Y0=354
    large_coords = {
        "management_system": fitz.Rect(87.9, 185, 580, 226.6),  # Same as standard
        "Company Name and Address": fitz.Rect(87.9, 229, 580, 295),  # Fixed: Match soft copy coordinates
        "ISO Standard": fitz.Rect(194.9, 300, 460.3, 336),  # Same as standard
        "Scope": fitz.Rect(85, 354, 577, 536),  # Y0=354 for all generation types
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # Logo template coordinates (when Logo = "yes")
    logo_coords = {
        "management_system": fitz.Rect(87.9, 175, 580, 216.6),  # Shifted up by 10pt from standard
        "logo": fitz.Rect(87.9, 206.6, 580, 242.6),  # Logo area: below management_system, above company name
        "Company Name and Address": fitz.Rect(87.9, 262.6, 580, 355),  # Lowered y-coordinates
        "ISO Standard": fitz.Rect(194.9, 334, 460.3, 370),  # Same as standard (not lowered)
        "Scope": {
            "short": fitz.Rect(87.9, 386, 580, 475),    # Updated scope coordinates for logo template
            "long": fitz.Rect(87.9, 373, 580, 486)      # Updated scope coordinates for logo template
        },
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # ✅ ADDED: Defensive checks for coordinate dictionaries
    if standard_coords is None:
        print(f"⚠️ [CERTIFICATE] standard_coords is None - cannot continue")
        raise ValueError("standard_coords is None - cannot generate certificate")
    
    if large_coords is None:
        print(f"⚠️ [CERTIFICATE] large_coords is None - cannot continue")
        raise ValueError("large_coords is None - cannot generate certificate")
    
    if logo_coords is None:
        print(f"⚠️ [CERTIFICATE] logo_coords is None - cannot continue")
        raise ValueError("logo_coords is None - cannot generate certificate")
    
    # Select coordinates based on template type
    if template_type in ["standard", "standard_eco", "standard_nonaccredited"]:
        coords = standard_coords
    elif template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other"]:
        coords = large_coords
    elif template_type in ["logo", "logo_nonaccredited", "logo_other", "logo_nonaccredited_other"]:
        coords = logo_coords
    else:
        # Default to standard coordinates
        coords = standard_coords
    
    # ✅ ADDED: Defensive check for selected coords
    if coords is None:
        raise ValueError("Selected coords is None - cannot generate certificate")

    
    font_starts = {
        "Company Name and Address": 45,  # Company Name starts from 45pt
        "Scope": 20,
        "ISO Standard": 80,
        "management_system": 15,  # Management system line font size
    }
    # --- End Configuration ---

    # Determine Scope coordinates based on content length
    scope_text = values.get("Scope", "")
    scope_words = len(scope_text.split())
    
    # Calculate estimated lines for Scope (approximate calculation)
    estimated_lines = max(1, (scope_words * 8) // 60)  # Rough estimate: 8 chars per word, 60 chars per line
    
    # ✅ ADDED: Defensive check for Scope coordinates
    if "Scope" not in coords:
        print(f"⚠️ [CERTIFICATE] Scope coordinates not found in coords - cannot continue")
        raise ValueError("Scope coordinates not found - cannot generate certificate")
    
    # Determine which coordinate set to use (lines win over words)
    if template_type in ["standard", "standard_eco", "standard_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_nonaccredited_other"]:
        # Standard template: dynamic coordinates based on content length
        if estimated_lines >= 24:  # Long content condition
            if "long" not in coords["Scope"]:
                print(f"⚠️ [CERTIFICATE] Scope long coordinates not found - cannot continue")
                raise ValueError("Scope long coordinates not found - cannot generate certificate")
            scope_rect = coords["Scope"]["long"]
            scope_layout = "long"
            print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines (≥24) -> selected LONG scope coordinates")

        else:  # Short content condition
            if "short" not in coords["Scope"]:
                print(f"⚠️ [CERTIFICATE] Scope short coordinates not found - cannot continue")
                raise ValueError("Scope short coordinates not found - cannot generate certificate")
            scope_rect = coords["Scope"]["short"]
            scope_layout = "short"
            print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines (<24) -> selected SHORT scope coordinates")

    else:
        # Large template: fixed large coordinates
        scope_rect = coords["Scope"]
        scope_layout = "large"
        print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines -> selected LARGE scope coordinates (fixed)")
    
    print(f"🎯 [CERTIFICATE] Final scope coordinates: {scope_rect}")
    print(f"[CHECK] template_type={template_type}, estimated_lines={estimated_lines}")
    print(f"[CHECK] coords['Scope'] (pre-set) type={type(coords['Scope']).__name__}, value={coords['Scope']}")
    print(f"[CHECK] scope_rect selected type={type(scope_rect).__name__}, value={scope_rect}")

    # Store original scope coordinates before modification (for Extra Line processing)
    original_scope_coords = coords["Scope"].copy() if isinstance(coords["Scope"], dict) else coords["Scope"]
    print(f"[CHECK] original_scope_coords type={type(original_scope_coords).__name__}")
    
    # Add Scope coordinates to the main coords dictionary
    coords["Scope"] = scope_rect
    print(f"[CHECK] coords['Scope'] (post-set) type={type(coords['Scope']).__name__}, value={coords['Scope']}")
    
    # ✅ NEW: Adjust scope coordinates when Extra Line is present with dynamic height logic
    extra_line = values.get("Extra Line", "").strip()
    if extra_line:
        print(f"🔍 [CERTIFICATE] Extra Line present - using dynamic scope height based on content length")
        
        # Calculate content length to determine appropriate scope height
        scope_text = values.get("Scope", "")
        scope_words = len(scope_text.split())
        estimated_lines = max(1, (scope_words * 8) // 60)  # Same calculation as standard templates
        
        if estimated_lines < 24:
            # Short scope: 89pt height (same as standard short scope)
            scope_rect = fitz.Rect(87.9, 386, 580, 475)  # Height: 89pt
            print(f"🔍 [CERTIFICATE] Extra Line - Short scope: {estimated_lines} lines, 89pt height")
        elif estimated_lines <= 30:
            # Long scope: 113pt height (same as standard long scope)
            scope_rect = fitz.Rect(87.9, 373, 580, 486)  # Height: 113pt
            print(f"🔍 [CERTIFICATE] Extra Line - Long scope: {estimated_lines} lines, 113pt height")
        else:
            # Large scope: 182pt height for >30 lines (same as large template)
            scope_rect = fitz.Rect(85, 354, 577, 536)    # Height: 182pt
            print(f"🔍 [CERTIFICATE] Extra Line - Large scope: {estimated_lines} lines, 182pt height")
        
        # Update the scope coordinates with dynamic height
        coords["Scope"] = scope_rect
        print(f"🔍 [CERTIFICATE] Extra Line scope coordinates set to: {scope_rect}")
        
    else:
        print(f"🔍 [CERTIFICATE] No Extra Line - using standard scope coordinates")
    
    # Function to insert logo with smart positioning
    def insert_logo_with_smart_positioning(page, logo_image, logo_rect):
        """
        Smart logo insertion that handles different aspect ratios:
        - Square logos: Use full coordinates, maintain aspect ratio
        - Horizontal logos: Center horizontally, maintain aspect ratio
        - Vertical logos: Center vertically, maintain aspect ratio
        """
        # Get logo dimensions
        logo_width = logo_image.width
        logo_height = logo_image.height
        
        # Calculate aspect ratios
        logo_aspect = logo_width / logo_height
        rect_aspect = logo_rect.width / logo_rect.height
        
        # Determine positioning strategy based on aspect ratio
        if logo_aspect > rect_aspect:
            # Logo is more horizontal than rectangle
            # Scale to fit width, center vertically
            scale_factor = logo_rect.width / logo_width
            new_width = logo_rect.width
            new_height = logo_height * scale_factor
            
            # Center vertically
            x = logo_rect.x0
            y = logo_rect.y0 + (logo_rect.height - new_height) / 2
            
        elif logo_aspect < rect_aspect:
            # Logo is more vertical than rectangle
            # Scale to fit height, center horizontally
            scale_factor = logo_rect.height / logo_height
            new_width = logo_width * scale_factor
            new_height = logo_rect.height
            
            # Center horizontally
            x = logo_rect.x0 + (logo_rect.width - new_width) / 2
            y = logo_rect.y0
            
        else:
            # Logo and rectangle have similar aspect ratios
            # Scale to fit within rectangle, center both ways
            scale_factor = min(logo_rect.width / logo_width, logo_rect.height / logo_height)
            new_width = logo_width * scale_factor
            new_height = logo_height * scale_factor
            
            # Center both horizontally and vertically
            x = logo_rect.x0 + (logo_rect.width - new_width) / 2
            y = logo_rect.y0 + (logo_rect.height - new_height) / 2
        
        # Convert PIL Image to PyMuPDF image
        try:
            # Convert PIL Image to bytes
            import io
            img_buffer = io.BytesIO()
            logo_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            img_bytes = img_buffer.read()
            
            # Insert into PDF
            page.insert_image(fitz.Rect(x, y, x + new_width, y + new_height), stream=img_bytes)
            print(f"🔍 [LOGO] Logo inserted with smart positioning: {new_width:.1f}x{new_height:.1f}")
            
        except Exception as e:
            print(f"❌ [LOGO] Error inserting logo: {e}")

    # ✅ ADDED: Shared logo functions for better logo handling
    def insert_logo_into_pdf(page, logo_file, logo_rect):
        """
        Insert logo into PDF with smart positioning
        """
        try:
            # Convert logo file to image
            logo_image = convert_file_to_image(logo_file)
            # Use smart positioning logic
            insert_logo_with_smart_positioning(page, logo_image, logo_rect)
            print(f"✅ [LOGO] Logo inserted successfully: {logo_file.name if hasattr(logo_file, 'name') else 'unknown'}")
        except Exception as e:
            print(f"❌ [LOGO] Failed to insert logo: {e}")

    def convert_file_to_image(file):
        """
        Convert uploaded file to PIL Image
        """
        try:
            from PIL import Image
            import io
            
            if hasattr(file, 'file'):
                # Reset file pointer
                file.file.seek(0)
                # Read file content
                file_content = file.file.read()
                # Convert to PIL Image
                logo_image = Image.open(io.BytesIO(file_content))
                return logo_image
            else:
                raise ValueError("File object has no file attribute")
        except Exception as e:
            print(f"❌ [LOGO] Error converting file to image: {e}")
            raise

    # ✅ ADDED: Render optional fields function
    def render_optional_fields(page, values, key_coords, value_coords, font_settings):
        """
        Render optional fields with dynamic positioning based on available data.
        """
        # ✅ ADDED: Defensive checks for None values
        if values is None:
            print(f"⚠️ [CERTIFICATE] Values dictionary is None in render_optional_fields - skipping")
            return
        
        if key_coords is None or value_coords is None:
            print(f"⚠️ [CERTIFICATE] Coordinates are None in render_optional_fields - skipping")
            return
        
        if font_settings is None:
            print(f"⚠️ [CERTIFICATE] Font settings is None in render_optional_fields - skipping")
            return
        # Define field order (top to bottom) and their display labels
        fields = [
            "Certificate Number",        # seq 1 - TOP (always present)
            "Initial Registration Date", # seq 2 (optional - not always present)
            "Original Issue Date",       # seq 3 (always present)
            "Issue Date",               # seq 4 (always present)
            "Surveillance Group",       # seq 5 (only 1 of 3 fields present)
            "Recertification Date"      # seq 6 - BOTTOM (always present)
        ]
        
        # Define the surveillance group fields (only 1 will be present)
        surveillance_group_fields = [
            "Surveillance/ Expiry Date",
            "Surveillance Due Date", 
            "Expiry Date"
        ]
        
        # ✅ ADDED: Spanish translation maps for field labels
        spanish_field_labels = {
            "Certificate Number": "Certificado No.",
            "Initial Registration Date": "Fecha de Registro Inicial",
            "Original Issue Date": "Fecha de Emisión Original",
            "Issue Date": "Fecha de Asunto",
            "Surveillance/ Expiry Date": "Vigilancia/Fecha de Caducidad",
            "Surveillance Due Date": "Fecha de Vigilancia Debida",
            "Expiry Date": "Fecha de Caducidad",
            "Recertification Date": "Fecha de Recertificación"
        }
        
        # English display labels for PDF rendering (cleaner, shorter)
        english_field_labels = {
            "Certificate Number": "Certificate No.",
            "Initial Registration Date": "Initial Registration Date",
            "Original Issue Date": "Original Issue Date",
            "Issue Date": "Issue Date", 
            "Surveillance/ Expiry Date": "Surveillance/ Expiry Date",
            "Surveillance Due Date": "Surveillance Due Date",
            "Expiry Date": "Expiry Date",
            "Recertification Date": "Recertification Date"
        }
        
        # ✅ ADDED: Select display labels based on language
        language = values.get("Language", "").strip().lower()
        if language == "s":
            display_labels = spanish_field_labels
            print(f"🔍 [CERTIFICATE] Using Spanish field labels")
        else:
            display_labels = english_field_labels
            print(f"🔍 [CERTIFICATE] Using English field labels")

        # Filter available fields (non-empty) with special handling for surveillance group
        available_fields = []
        
        # ✅ ADDED: Debug logging for values received
        print(f"🔍 [CERTIFICATE] ===== OPTIONAL FIELDS DEBUG =====")
        print(f"🔍 [CERTIFICATE] Values received: {list(values.keys())}")
        print(f"🔍 [CERTIFICATE] Surveillance/ Expiry Date: '{values.get('Surveillance/ Expiry Date', '')}'")
        print(f"🔍 [CERTIFICATE] Surveillance Due Date: '{values.get('Surveillance Due Date', '')}'")
        print(f"🔍 [CERTIFICATE] Expiry Date: '{values.get('Expiry Date', '')}'")
        
        for field in fields:
            if field == "Surveillance Group":
                # Handle surveillance group - find which field is present
                surveillance_value = None
                surveillance_label = None
                
                print(f"🔍 [CERTIFICATE] Processing Surveillance Group...")
                for surveillance_field in surveillance_group_fields:
                    print(f"🔍 [CERTIFICATE] Checking '{surveillance_field}': '{values.get(surveillance_field, '')}'")
                    if surveillance_field in values and values[surveillance_field]:
                        surveillance_value = values[surveillance_field]
                        surveillance_label = surveillance_field
                        print(f"🔍 [CERTIFICATE] Found surveillance field: '{surveillance_label}' = '{surveillance_value}'")
                        break
                
                if surveillance_value and surveillance_label:
                    available_fields.append((surveillance_label, surveillance_value))
                    print(f"🔍 [CERTIFICATE] Added surveillance field to available fields")
                else:
                    print(f"🔍 [CERTIFICATE] No surveillance field found or all are empty")
            else:
                value = values.get(field, "").strip()
                if value:  # Only include non-empty fields
                    available_fields.append((field, value))
                    print(f"🔍 [CERTIFICATE] Added field '{field}' = '{value}' to available fields")
                else:
                    print(f"🔍 [CERTIFICATE] Field '{field}' is empty, skipping")
        
        print(f"🔍 [CERTIFICATE] Total available fields: {len(available_fields)}")
        print(f"🔍 [CERTIFICATE] Available fields: {available_fields}")
        print(f"🔍 [CERTIFICATE] ===== END OPTIONAL FIELDS DEBUG =====")

        # Calculate starting position using formula: (6 - available_count) + 1
        total_fields = 6
        available_count = len(available_fields)
        starting_position = (6 - available_count) + 1

        # Essential logging only
        print(f"🔍 [CERTIFICATE] Optional fields: {available_count}/{total_fields} available")

        # Render from starting position
        for i, (field, value) in enumerate(available_fields):
            coord_index = starting_position - 1 + i  # Convert to 0-based index

            if coord_index < len(key_coords):  # Prevent overflow
                # Prepare text with custom display labels
                key_text = display_labels[field]  # Use custom display label
                value_text = f":{value}"            # Colon + value (no space)

                # Essential field logging
                print(f"🔍 [CERTIFICATE] Rendering: {field}")

                # Extract (x, y) coordinates from rectangles
                key_x = key_coords[coord_index].x0
                key_y = key_coords[coord_index].y0
                value_x = value_coords[coord_index].x0
                value_y = value_coords[coord_index].y0

                # ✅ ADDED: Defensive checks for font settings
                fontsize = font_settings.get('fontsize', 15)  # Default to 15pt
                fontname = font_settings.get('fontname', 'Times-Roman')  # Default to Times-Roman
                color = font_settings.get('color', (0, 0, 0))  # Default to black
                
                # Insert at respective coordinates using (x, y) points
                safe_insert_text(
                    page,
                    (key_x, key_y),
                    key_text,
                    fontsize=fontsize,
                    fontname=fontname,
                    color=color
                )

                safe_insert_text(
                    page,
                    (value_x, value_y),
                    value_text,
                    fontsize=fontsize,
                    fontname=fontname,
                    color=color
                )

                print(f"   ✅ Rendered successfully at position {coord_index + 1}")
            else:
                print(f"⚠️ [CERTIFICATE] Warning: Coordinate index {coord_index} out of bounds")

        print(f"🔍 [CERTIFICATE] ===== END OPTIONAL FIELDS ANALYSIS =====\n")

    # ✅ ADDED: Optional field coordinates for large template
    large_optional_key_coordinates = [
        fitz.Rect(175.5, 522, 343, 530),    # Row 1: Certificate Number
        fitz.Rect(175.5, 538, 343, 548),    # Row 2: Initial Registration Date
        fitz.Rect(175.5, 556, 343, 566),    # Row 3: Original Issue Date
        fitz.Rect(175.5, 574, 343, 584),    # Row 4: Issue Date
        fitz.Rect(175.5, 592, 343, 602),    # Row 5: Surveillance Group (only 1 field)
        fitz.Rect(175.5, 610, 343, 620)     # Row 6: Recertification Date
    ]
    
    large_optional_value_coordinates = [
        fitz.Rect(362.1, 522, 446.4, 530),  # Row 1: Certificate Number value
        fitz.Rect(362.1, 538, 446.4, 548),  # Row 2: Initial Registration Date value
        fitz.Rect(362.1, 556, 446.4, 566),  # Row 3: Original Issue Date value
        fitz.Rect(362.1, 574, 446.4, 584),  # Row 4: Issue Date value
        fitz.Rect(362.1, 592, 446.4, 602),  # Row 5: Surveillance Group value (only 1 field)
        fitz.Rect(362.1, 610, 446.4, 620)   # Row 6: Recertification Date value
    ]
    
    # ✅ ADDED: Optional field coordinates for standard template
    standard_optional_key_coordinates = [
        fitz.Rect(175.5, 499.1, 343, 509.1),    # Row 1: Certificate Number
        fitz.Rect(175.5, 516.9, 343, 526.9),    # Row 2: Initial Registration Date
        fitz.Rect(175.5, 535.1, 343, 545.1),    # Row 3: Original Issue Date
        fitz.Rect(175.5, 553.9, 343, 563.9),    # Row 4: Issue Date
        fitz.Rect(175.5, 571.6, 343, 581.6),    # Row 5: Surveillance Group (only 1 field)
        fitz.Rect(175.5, 589.3, 343, 599.3)     # Row 6: Recertification Date
    ]
    
    standard_optional_value_coordinates = [
        fitz.Rect(362.1, 499.1, 446.4, 509.1),  # Row 1: Certificate Number value
        fitz.Rect(362.1, 516.9, 446.4, 526.9),  # Row 2: Initial Registration Date value
        fitz.Rect(362.1, 535.1, 446.4, 545.1),  # Row 3: Original Issue Date value
        fitz.Rect(362.1, 553.9, 446.4, 563.9),  # Row 4: Issue Date value
        fitz.Rect(362.1, 571.6, 446.4, 581.6),  # Row 5: Surveillance Group value (only 1 field)
        fitz.Rect(362.1, 589.3, 446.4, 599.3)   # Row 6: Recertification Date value
    ]
    
    # ✅ UPDATED: Font settings for optional fields (matching soft copy)
    optional_font_settings = {
        "fontname": "Times-Roman",  # Same as soft copy
        "fontsize": 13,             # Same as soft copy (was 15pt)
        "color": (0, 0, 0)         # Black
    }

    # ✅ ADDED: Adjust scope coordinates based on whether Initial Registration Date is present
    initial_registration_date = values.get("Initial Registration Date", "")
    if initial_registration_date and initial_registration_date.strip():
        # When Initial Registration Date is present, reduce scope height to accommodate the extra field
        print(f"🔍 [CERTIFICATE] Initial Registration Date present - adjusting scope coordinates for large template")
        if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other"]:
            # Adjust large template scope coordinates
            coords["Scope"] = fitz.Rect(85, 351, 577, 520)  # Reduced height by 16 units
            print(f"🔍 [CERTIFICATE] Adjusted large template scope coordinates for Initial Registration Date")
    else:
        print(f"🔍 [CERTIFICATE] Using standard scope coordinates (Initial Registration Date not present)")

    # Process optional fields
    if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other"]:
        render_optional_fields(page, values, large_optional_key_coordinates, large_optional_value_coordinates, optional_font_settings)
    else:
        render_optional_fields(page, values, standard_optional_key_coordinates, standard_optional_value_coordinates, optional_font_settings)

    # Calculate optional fields count for field processing
    optional_fields_count = 0
    # ✅ ADDED: Defensive check for values before optional fields calculation
    if values is not None:
        for field in ["Certificate Number", "Initial Registration Date", "Original Issue Date", "Issue Date", "Surveillance/ Expiry Date", "Surveillance Due Date", "Expiry Date", "Recertification Date"]:
            if values.get(field, "").strip():
                optional_fields_count += 1
    else:
        print(f"⚠️ [CERTIFICATE] Values is None during optional fields calculation - using 0")
    
    # Process each field with enhanced text handling
    # ✅ FIELD CLASSIFICATION SYSTEM:
    # ==========================================
    # RENDERING_FIELDS: Main content fields rendered using main coordinates
    # OPTIONAL_FIELDS: Additional fields rendered using separate optional coordinates  
    # METADATA_FIELDS: Template selection fields (NOT rendered on PDF)
    # ==========================================
    RENDERING_FIELDS = ["Company Name", "Scope", "ISO Standard", "management_system", "Extra Line"]
    OPTIONAL_FIELDS = ["Certificate Number", "Initial Registration Date", "Original Issue Date", "Issue Date", "Surveillance/ Expiry Date", "Surveillance Due Date", "Expiry Date", "Recertification Date"]
    METADATA_FIELDS = ["Size", "Accreditation", "Logo", "Country", "logo_lookup"]
    
    # ✅ FIELD PROCESSING FLOW:
    # 1. Metadata fields (Size, Accreditation, Logo, Country) -> SKIPPED (no rendering needed)
    # 2. Optional fields (Certificate Number, Dates, etc.) -> SKIPPED (rendered via render_optional_fields)
    # 3. Main rendering fields (Company Name, Scope, ISO Standard, management_system) -> PROCESSED (main coordinates)
    
    # ✅ ADDED: Defensive check for values before main field processing
    if values is None:
        print(f"⚠️ [CERTIFICATE] Values dictionary is None in main field processing - cannot continue")
        raise ValueError("Values dictionary is None - cannot generate certificate")
    
    # Scope text now uses justification (left and right alignment) for professional appearance
    for field, text in values.items():
        # ✅ ADDED: Skip metadata fields that don't need rendering
        if field in METADATA_FIELDS:
            continue
            
        # ✅ ADDED: Skip optional fields as they are handled separately
        if field in OPTIONAL_FIELDS:
            continue
        
        if field == "Company Name":
            company_text = str(text) if text is not None else ""
            
            # Handle Company Name and Address together
            address_text = values.get("Address", "")
            safe_address_text = str(address_text) if address_text is not None else ""
            
            # ✅ ADDED: Extract Excel adjustment and font size values
            name_adjustment = parse_excel_adjustment(values.get("Name Adjustment", ""))
            name_font_size_adjustment = parse_excel_font_size(values.get("Name Font Size", ""))
            address_adjustment = parse_excel_adjustment(values.get("Address Adjustment", ""))
            address_font_size_adjustment = parse_excel_font_size(values.get("Address Font Size", ""))


            
            # PRE-PROCESS: Apply line break logic BEFORE font size calculation
            # This ensures both font calculation and rendering use the same processed text
            
            # Process Company Name with line break preservation
            def process_text_with_line_breaks(input_text, field_name):
                if not input_text:
                    return []
                
                # ✅ UPDATED: Split by actual line breaks and PRESERVE empty lines
                lines = input_text.split('\n')
                processed_lines = []
                
                for line in lines:
                    # Preserve empty lines to maintain spacing from Excel
                    processed_lines.append(line)  # Keep original line (including empty ones)

                
                return processed_lines
            
            # Pre-process both company and address text to handle line breaks
            company_processed_lines = process_text_with_line_breaks(company_text, "Company")
            address_processed_lines = process_text_with_line_breaks(safe_address_text, "Address")
            
            # ✅ ADDED: Determine address alignment based on Excel column or line count
            address_alignment_column = values.get("Address alignment", "").strip().lower()
            address_lines_count = len(address_processed_lines)
            
            if address_alignment_column == "center":
                address_alignment = "center"
                print(f"🔍 [CERTIFICATE] Address: Excel column specifies CENTERED alignment")
            elif address_alignment_column == "left":
                address_alignment = "left"
                print(f"🔍 [CERTIFICATE] Address: Excel column specifies LEFT alignment")
            else:
                # Default logic: always center unless Excel column specifies otherwise
                address_alignment = "center"  # Default: always center
                print(f"🔍 [CERTIFICATE] Address: No Excel column value - using CENTERED alignment (default)")
            
            # ✅ ADDED: Defensive check for coords dictionary
            if coords is None:
                print(f"⚠️ [CERTIFICATE] Coords dictionary is None - cannot continue")
                raise ValueError("Coords dictionary is None - cannot generate certificate")
            
            if "Company Name and Address" not in coords:
                print(f"⚠️ [CERTIFICATE] Company Name and Address coordinates not found - cannot continue")
                raise ValueError("Company Name and Address coordinates not found - cannot generate certificate")
            
            rect = coords["Company Name and Address"]
            
            # Check if address text is empty or None
            if not safe_address_text:
                print(f"[WARNING] [COMPANY ADDRESS] WARNING: Address text is empty or None!")
            elif safe_address_text.strip() == "":
                print(f"[WARNING] [COMPANY ADDRESS] WARNING: Address text is only whitespace!")
            else:
                print(f"[SUCCESS] [COMPANY ADDRESS] Address text is valid and non-empty")
            
            # Combine text with natural spacing (single line break)
            combined_text = f"{company_text}\n{safe_address_text}"  # \n creates ~2-3pt spacing
            
            # ✅ ADDED: Defensive check for font_starts dictionary
            if font_starts is None:
                print(f"⚠️ [CERTIFICATE] Font_starts dictionary is None - using default")
                start_size = 30
            else:
                start_size = font_starts.get("Company Name and Address", 30)
            font_size = start_size
            
    
            

            
            # ✅ UPDATED: Dynamic Company Name font sizing based on line count
            # First, determine if Company Name will be single line or multi-line
            company_lines_count = len([line for line in company_processed_lines if line.strip()])
            
            # Set initial font size based on line count
            if company_lines_count <= 1:
                company_font_size = 35  # Single line - start with 35pt
                print(f"🔍 [CERTIFICATE] Company Name: Single line detected, starting with {company_font_size}pt")
            else:
                company_font_size = 30  # Multiple lines - start with 30pt
                print(f"🔍 [CERTIFICATE] Company Name: {company_lines_count} lines detected, starting with {company_font_size}pt")
            
            address_font_size = 13.6
            
            # Variables to store the final wrapped lines and font sizes
            final_company_lines = []
            final_address_lines = []
            
            # ✅ IMPROVED: Different logic for single line vs multi-line company names
            if company_lines_count <= 1:
                # NO cmd+enter in Excel: Force single line, use font reduction only
                print(f"🔍 [CERTIFICATE] No cmd+enter detected - forcing single line with font reduction")
                
                while company_font_size >= 8:  # Minimum font size
                    # Check if entire company name fits in one line at current font size
                    font_obj = fitz.Font(fontname=fontname)
                    text_width = font_obj.text_length(company_text, company_font_size)
                    
                    if text_width <= rect.width - 10:  # Leave margin
                        # Text fits in one line - use this font size
                        final_company_lines = [company_text]  # Single line
                        print(f"✅ [CERTIFICATE] Company name fits in one line at {company_font_size}pt (width: {text_width:.1f}pt)")
                        break
                    else:
                        # Text too wide - reduce font size and try again
                        print(f"🔍 [CERTIFICATE] Company name too wide at {company_font_size}pt (width: {text_width:.1f}pt > {rect.width - 10:.1f}pt), reducing to {company_font_size - 1}pt")
                        company_font_size -= 1
                
                # If we reached minimum font size and still doesn't fit, use the minimum
                if company_font_size < 8:
                    company_font_size = 8
                    final_company_lines = [company_text]
                    print(f"⚠️ [CERTIFICATE] Company name forced to minimum font size 8pt")
                
            else:
                # cmd+enter present in Excel: Allow word wrapping up to 2 lines
                print(f"🔍 [CERTIFICATE] cmd+enter detected - allowing word wrapping up to 2 lines")
                
            while company_font_size >= 8:  # Minimum font size
                company_lines = []
                
                # Process each pre-processed line with word wrapping
                for processed_line in company_processed_lines:
                    if not processed_line.strip():  # Empty line - preserve it
                        company_lines.append("")  # Add empty line to maintain spacing
                        continue
                    
                    # Non-empty line - apply word wrapping
                    words = processed_line.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        if font_obj.text_length(test_line, company_font_size) <= rect.width - 10:  # Leave margin
                            current_line = test_line
                        else:
                            if current_line:
                                company_lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        company_lines.append(current_line)
                
                # ✅ UPDATED: Allow Company Name to use up to 2 lines (after line breaks + word wrapping)
                if len(company_lines) <= 2:
                    final_company_lines = company_lines.copy()
                    break
                
                # Reduce Company Name font size
                company_font_size -= 1
            
            # ✅ UPDATED: Apply Excel font size adjustment AFTER optimization (user override)
            if name_font_size_adjustment != 0:
                optimized_company_font = company_font_size
                company_font_size += name_font_size_adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Company Name Font Size adjustment AFTER optimization: {optimized_company_font}pt + {name_font_size_adjustment}pt = {company_font_size}pt")
            
            # ✅ PHASE 0: Address Line Count-Based Height Allocation
            # Determine if company name is single-line or multi-line
            has_multiple_company_lines = len(final_company_lines) > 1
            address_lines_count = len(address_processed_lines)
            
            if has_multiple_company_lines:
                # Multi-line company: Use fixed height allocation based on template and address lines
                if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "large_nonaccredited_other"]:
                    if address_lines_count == 1:
                        company_height = 42  # Name +8, Address -5
                    elif address_lines_count == 2:
                        company_height = 33  # Name +8, Address -5 (was 25pt - FIXED)
                    else:  # 3+ lines
                        company_height = 19  # Name 19pt, Address 37pt (fits in 56pt total)
                elif template_type in ["standard", "standard_eco", "standard_nonaccredited", "standard_other", "standard_other_eco", "standard_nonaccredited_other"]:
                    if address_lines_count == 1:
                        company_height = 45  # Match single-line allocation
                    else:
                        company_height = 30  # Less space when address is multi-line
                else:
                    # Logo template - same height allocation as large template
                    if address_lines_count == 1:
                        company_height = 42  # Name +8, Address -5 (same as large)
                    elif address_lines_count == 2:
                        company_height = 33  # Name +8, Address -5 (same as large)
                    else:  # 3+ lines
                        company_height = 19  # Name 19pt, Address 37pt (same as large)
            else:
                # Single line company: dynamic height based on template and address lines
                if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "large_nonaccredited_other"]:
                    if address_lines_count == 1:
                        company_height = 42  # Match softCopy allocation
                    elif address_lines_count == 2:
                        company_height = 33  # Same as multi-line for 2 address lines
                    else:  # 3+ lines
                        company_height = 19  # Match softCopy allocation
                elif template_type in ["standard", "standard_eco", "standard_nonaccredited", "standard_other", "standard_other_eco", "standard_nonaccredited_other"]:
                    if address_lines_count == 1:
                        company_height = 45
                    else:
                        company_height = 30
                else:
                    # Logo template - same height allocation as large template
                    if address_lines_count == 1:
                        company_height = 42  # Name +8, Address -5 (same as large)
                    elif address_lines_count == 2:
                        company_height = 33  # Name +8, Address -5 (same as large)
                    else:  # 3+ lines
                        company_height = 19  # Name 19pt, Address 37pt (same as large)
            
            print(f"🔍 [CERTIFICATE DEBUG] Company height allocation: {company_height}pt (company lines: {len(final_company_lines)}, address lines: {address_lines_count}, template: {template_type})")
            
            # ✅ ENHANCED: Height-Aware Font Reduction for Company Name
            # Apply the same sophisticated font reduction logic as soft copy
            if len(final_company_lines) > 1:
                # Multi-line: Calculate required font size to fit allocated height
                # Each line needs: font_size * 1.002 (0.1pt spacing)
                # Total height = font_size * 1.002 * number_of_lines
                # So: font_size = company_height / (1.002 * number_of_lines)
                required_font_size = company_height / (1.002 * len(final_company_lines))
                if required_font_size < company_font_size:
                    company_font_size = required_font_size
                    print(f"🔍 [CERTIFICATE DEBUG] Multi-line font reduced to {company_font_size:.1f}pt to fit {company_height}pt height")
            else:
                # Single line: Calculate required font size to fit allocated height
                # Single line needs: font_size * 1.0 (no spacing)
                # So: font_size = company_height
                required_font_size = company_height
                if required_font_size < company_font_size:
                    company_font_size = required_font_size
                    print(f"🔍 [CERTIFICATE DEBUG] Single line font reduced to {company_font_size:.1f}pt to fit {company_height}pt height")
            
            # Ensure minimum font size
            if company_font_size < 8:
                company_font_size = 8
                print(f"🔍 [CERTIFICATE DEBUG] Company font size set to minimum 8pt")
            
            # ✅ REMOVED: Name Font Size adjustment now applied BEFORE optimization (see above)
            
            # ✅ PHASE 2: Address Positioning Optimization
            # Now find font size for Address to fit in remaining space
            remaining_height = rect.height - company_height  # No margin - address starts immediately

            
            address_font_size_attempts = 0
            while address_font_size >= 6:  # Minimum font size
                address_font_size_attempts += 1

                
                # Process Address using pre-processed lines with word wrapping
                address_lines = []
                max_address_width = rect.width - 10  # Leave margin
                
                # Process each pre-processed address line with word wrapping
                for line_idx, processed_line in enumerate(address_processed_lines):
                    if not processed_line.strip():  # Empty line - preserve it
                        address_lines.append("")  # Add empty line to maintain spacing
                        continue
                    
                    # Non-empty line - apply word wrapping
                    words = processed_line.split()
                    current_line = ""
                    original_line_wrapped = False
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        test_width = font_obj.text_length(test_line, address_font_size)
                        if test_width <= max_address_width:  # Leave margin
                            current_line = test_line
                        else:
                            # Line too wide - wrap it
                            if current_line:
                                # Log the wrapped line before adding it
                                line_width = font_obj.text_length(current_line, address_font_size)
                                print(f"🔍 [ADDRESS WIDTH] Line wrapped at {address_font_size:.1f}pt: '{current_line[:50]}{'...' if len(current_line) > 50 else ''}' (width: {line_width:.1f}pt <= {max_address_width:.1f}pt)")
                                address_lines.append(current_line)
                                original_line_wrapped = True
                            # Check if single word itself is too wide
                            word_width = font_obj.text_length(word, address_font_size)
                            if word_width > max_address_width:
                                print(f"⚠️ [ADDRESS WIDTH] Single word exceeds width: '{word}' (width: {word_width:.1f}pt > {max_address_width:.1f}pt) - will be truncated")
                            current_line = word
                    
                    if current_line:
                        # Log final line of this processed segment only if it was wrapped or might overflow
                        line_width = font_obj.text_length(current_line, address_font_size)
                        if original_line_wrapped or line_width > max_address_width:
                            print(f"🔍 [ADDRESS WIDTH] Final line segment at {address_font_size:.1f}pt: '{current_line[:50]}{'...' if len(current_line) > 50 else ''}' (width: {line_width:.1f}pt)")
                        address_lines.append(current_line)
                
                # Calculate Address height
                # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                    address_height = len(address_lines) * address_font_size * 1.1  # Tight spacing for large/logo templates
                else:  # standard templates
                    address_height = len(address_lines) * address_font_size * 1.2  # Loose spacing for standard templates

                
                # Check if Address fits in remaining space
                if address_height <= remaining_height:
                    final_address_lines = address_lines.copy()
                    
                    # ✅ ADDED: Final width check for all address lines (only when solution found)
                    font_obj = fitz.Font(fontname=fontname)
                    overflow_detected = False
                    print(f"🔍 [ADDRESS WIDTH] Final width check at {address_font_size:.1f}pt (max width: {max_address_width:.1f}pt):")
                    for line_idx, line in enumerate(final_address_lines):
                        if line.strip():  # Only check non-empty lines
                            line_width = font_obj.text_length(line, address_font_size)
                            if line_width > max_address_width:
                                overflow_detected = True
                                print(f"  ❌ Line {line_idx + 1} exceeds: '{line[:50]}{'...' if len(line) > 50 else ''}' (width: {line_width:.1f}pt > {max_address_width:.1f}pt)")
                            else:
                                print(f"  ✅ Line {line_idx + 1} fits: '{line[:50]}{'...' if len(line) > 50 else ''}' (width: {line_width:.1f}pt <= {max_address_width:.1f}pt)")
                    
                    if overflow_detected:
                        print(f"⚠️ [ADDRESS WIDTH] ⚠️ WARNING: Some address lines exceed available width at {address_font_size:.1f}pt")
                    
                    print(f"[SUCCESS] [COMPANY ADDRESS] Address fits! Final font size: {address_font_size}pt")
                    break
                else:
                    print(f"[ERROR] [COMPANY ADDRESS] Address too tall: {address_height:.1f}pt > {remaining_height:.1f}pt, reducing font size")
                
                # Reduce Address font size
                address_font_size -= 0.5
            
            # ✅ UPDATED: Apply Excel font size adjustment AFTER optimization (user override)
            if address_font_size_adjustment != 0:
                optimized_address_font = address_font_size
                address_font_size += address_font_size_adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Address Font Size adjustment AFTER optimization: {optimized_address_font}pt + {address_font_size_adjustment}pt = {address_font_size}pt")
            
            # Now render Company Name and Address dynamically
            if final_company_lines or final_address_lines:
                


                
                # Calculate total height
                total_height = company_height + address_height

                
                # No top margin - start at exact rectangle top
                start_y = rect.y0 + name_adjustment  # Start at exact top of box + Excel adjustment

                
                # Render Company Name first (starts from top)
                current_y = start_y
                for i, line in enumerate(final_company_lines):
                    # Consistent line spacing: 1.02 for all templates
                    line_height = company_font_size * 1.02  # Consistent spacing for all templates
                    
                    # ✅ PHASE 1: Company Name Baseline Positioning Fix
                    # Apply baseline offset only for single-line company names
                    if not has_multiple_company_lines:
                        # Single-line: Apply baseline offset to prevent overlap
                        y_pos = current_y + company_font_size * 0.2  # Baseline adjustment for single line
                    else:
                        # Multi-line: No offset to maintain proper line spacing
                        y_pos = current_y + 0  # No adjustment for multi-line
                    center_x = (rect.x0 + rect.x1) / 2
                    
                                        # ✅ UPDATED: Handle empty lines (preserve spacing from Excel)
                    if line.strip():  # Non-empty line - render text
                        # ✅ ENHANCED: Use mixed format text rendering for bold detection
                        if '**' in line or '__' in line:
                            # Calculate total width for centering
                            segments = process_bold_text(line)
                            total_width = 0
                            for segment_text, _, _ in segments:
                                if segment_text:
                                    font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                    total_width += font_obj.text_length(segment_text, company_font_size)
                            
                            x_pos = center_x - total_width / 2
                            render_mixed_format_text(page, (x_pos, y_pos), line, company_font_size, color)
                        else:
                            # Standard rendering for non-bold text
                            line_width = font_obj.text_length(line, company_font_size)
                            x_pos = center_x - line_width / 2
                            
                            # ✅ ADDED: Dynamic font selection for Company Name
                            dynamic_font = get_font_for_text(line, fontname)
                            safe_insert_text(
                                page,
                                (x_pos, y_pos),
                                line,
                                fontsize=company_font_size,
                                fontname=dynamic_font,
                                color=color
                            )
                    # Empty line - just advance position (creates blank space)
                    
                    # ✅ PHASE 1: Company Block Advance Compensation
                    # For single-line company names, subtract the applied baseline offset to avoid introducing extra gap
                    if not has_multiple_company_lines:
                        current_y += line_height - (company_font_size * 0.2)
                    else:
                        current_y += line_height
                
                # ✅ PHASE 2: Address Positioning Optimization
                # No spacing - address starts immediately after company name
                
                # Now render Address below Company Name (fills remaining space)
                has_multiple_address_lines = len(final_address_lines) > 1
                
                for i, line in enumerate(final_address_lines):
                    # Consistent line spacing: 1.05 for all templates
                    line_height = address_font_size * 1.05  # Consistent spacing for all templates
                    
                    # ✅ PHASE 3: Address Baseline Compensation System
                    # Apply compensation for single-line address regardless of company line count
                    if has_multiple_address_lines:
                        # Multi-line address: no baseline centering, no compensation
                        y_pos = current_y + address_adjustment  # Excel adjustment only
                    else:
                        # Single-line address: keep baseline centering with compensation
                        y_pos = current_y + (line_height / 2) - (company_font_size * 0.2) + address_adjustment  # Centered + compensation + Excel adjustment
                    
                    # ✅ ADDED: Dynamic alignment based on address line count
                    if address_alignment == "center":
                        center_x = (rect.x0 + rect.x1) / 2
                        line_width = font_obj.text_length(line, address_font_size)
                        x_pos = center_x - line_width / 2
                    else:  # left-aligned
                        x_pos = rect.x0 + 5  # Small left margin
                    
                    # ✅ UPDATED: Handle empty lines (preserve spacing from Excel)
                    if line.strip():  # Non-empty line - render text
                        # ✅ ENHANCED: Use mixed format text rendering for bold detection
                        if '**' in line or '__' in line:
                            # Calculate total width for alignment
                            segments = process_bold_text(line)
                            total_width = 0
                            for segment_text, _, _ in segments:
                                if segment_text:
                                    font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                    total_width += font_obj.text_length(segment_text, address_font_size)
                            
                            # Apply alignment based on address_alignment setting
                            if address_alignment == "center":
                                x_pos = center_x - total_width / 2
                            else:  # left-aligned
                                x_pos = rect.x0 + 5  # Small left margin
                            
                            render_mixed_format_text(page, (x_pos, y_pos), line, address_font_size, color)
                        else:
                            # Standard rendering for non-bold text
                            # ✅ FIXED: Don't recalculate x_pos - use the alignment already determined
                            
                            # ✅ ADDED: Dynamic font selection for Address
                            dynamic_font = get_font_for_text(line, fontname)
                            safe_insert_text(
                                page,
                                (x_pos, y_pos),
                                line,
                                fontsize=address_font_size,
                                fontname=dynamic_font,
                                color=color
                            )
                    # Empty line - just advance position (creates blank space)
                    
                    current_y += line_height
                
                # Print final font sizes for Company Name and Address

                
                # Skip Address processing since it's handled above
                continue
            
        elif field == "Address":
            # Skip Address processing since it's handled with Company Name
            continue
        
        elif field == "ISO Standard":
            # After processing ISO Standard, render the management system line
            iso_standard_text = text
            
            # Expand the ISO standard to full version with year
            expanded_iso = expand_iso_standard(iso_standard_text)
            
            # Get the language preference first
            language = values.get("Language", "").strip().lower()
            
            
            # Get the description from the appropriate language mapping
            if language == "s":
                system_name = ISO_STANDARDS_DESCRIPTIONS_SPANISH.get(expanded_iso, "Sistema de Gestión")
            else:
                system_name = ISO_STANDARDS_DESCRIPTIONS.get(expanded_iso, "Management System")
            
            # Capitalize first letters of each word in system_name, with special handling for acronyms
            def capitalize_management_system(name, is_spanish=False):
                words = name.split()
                result = []
                for word in words:
                    if is_spanish:
                        # Spanish acronyms and special cases
                        if word.upper() in ['IT', 'ISO', 'IEC', 'TI', 'SGC', 'SGA', 'SGSST', 'SGSE', 'SGSI', 'SGSA', 'SGAS']:
                            result.append(word.upper())
                        elif word.lower() in ['el', 'la', 'de', 'del', 'en', 'y', 'o', 'con', 'para', 'por']:
                            # Keep Spanish articles and prepositions lowercase
                            result.append(word.lower())
                        else:
                            # Capitalize first letter of each word
                            result.append(word.capitalize())
                    else:
                        # English acronyms
                        if word.upper() in ['IT', 'ISO', 'IEC', 'OH&S', 'HSE', 'EMS', 'QMS', 'FSMS', 'ISMS', 'ABMS']:
                            result.append(word.upper())
                        else:
                            result.append(word.capitalize())
                return ' '.join(result)
            
            system_name_caps = capitalize_management_system(system_name, is_spanish=(language == "s"))
            
            # Create the management system line with Language support
            if language == "s":
                management_line = f"Esto es para certificar que {system_name_caps} de"
            else:
                management_line = f"This is to certify that the {system_name_caps} of"
            
            # Get the management_system rectangle
            management_rect = coords["management_system"]
            
            # Calculate center position for the text
            center_x = (management_rect.x0 + management_rect.x1) / 2
            center_y = (management_rect.y0 + management_rect.y1) / 2 + 15/3  # Adjust for baseline
            
            # Determine initial font size: 12pt for Spanish + ISO 45001, otherwise 15pt
            management_font_size = 12 if (language == "s" and "45001" in expanded_iso) else 15
            
            # ✅ ADDED: Font size reduction logic to prevent x-coordinate overflow
            max_width = management_rect.width - 10  # Leave 5pt margin on each side (87.9 to 580 = 492.1pt width)
            print(f"🔍 [CERTIFICATE] Management line overflow protection: max_width={max_width:.1f}pt")
            
            while management_font_size >= 8:  # Minimum font size
                font_obj = fitz.Font(fontname="Times-BoldItalic")
                text_width = font_obj.text_length(management_line, management_font_size)
                
                if text_width <= max_width:
                    print(f"✅ [CERTIFICATE] Management line fits at {management_font_size}pt (width: {text_width:.1f}pt)")
                    break  # Text fits!
                
                print(f"🔍 [CERTIFICATE] Management line too wide at {management_font_size}pt (width: {text_width:.1f}pt > {max_width:.1f}pt), reducing to {management_font_size - 0.5}pt")
                management_font_size -= 0.5  # Reduce font size
            
            # Ensure minimum font size
            if management_font_size < 8:
                management_font_size = 8
                print(f"⚠️ [CERTIFICATE] Management line forced to minimum font size 8pt")
            
            # Calculate text width for centering with final font size
            font_obj = fitz.Font(fontname="Times-BoldItalic")  # Use bold italic font
            text_width = font_obj.text_length(management_line, management_font_size)
            start_x = center_x - text_width / 2
            
            # Insert the management system text
            safe_insert_text(
                page,
                (start_x, center_y),
                management_line,
                fontsize=management_font_size,
                fontname="Times-BoldItalic", # Bold italic font
                color=(0, 0, 0)  # Black color
            )
            
            # Print font size for management system

            
            # ✅ NEW: Render ISO Standard with centered positioning (like softcopy)
            # Get ISO Standard coordinates
            iso_rect = coords["ISO Standard"]
            print(f"🔍 [CERTIFICATE DEBUG] ISO Standard coordinates: {iso_rect}")
            
            # Use expanded ISO for display
            iso_text = expanded_iso
            iso_font_size = 30  # Standard font size for ISO Standard
            iso_fontname = "Times-Bold"
            iso_color = (0, 0, 0)
            
            # Perfect centering for ISO Standard - both horizontal and vertical (same as softcopy)
            center_x = (iso_rect.x0 + iso_rect.x1) / 2
            center_y = (iso_rect.y0 + iso_rect.y1) / 2 + iso_font_size/3  # Adjust for baseline
            
            # Calculate text width for centering
            font_obj = fitz.Font(fontname=iso_fontname)
            text_width = font_obj.text_length(iso_text, iso_font_size)
            start_x = center_x - text_width / 2
            
            # Insert the ISO Standard text with centered positioning
            safe_insert_text(
                page,
                (start_x, center_y),
                iso_text,
                fontsize=iso_font_size,
                fontname=iso_fontname,
                color=iso_color
            )
            
            print(f"🔍 [CERTIFICATE DEBUG] ISO Standard final rendering: text='{iso_text}', font_size={iso_font_size}pt, center=({center_x:.1f}, {center_y:.1f})")
            print(f"✅ [CERTIFICATE] ISO Standard field processed successfully")
            
            # Skip the normal field processing since we handled it above
            continue
        
        # ✅ VALIDATION: Ensure only rendering fields are processed
        if field not in RENDERING_FIELDS:
            continue
        
        # Handle other fields normally
        if field not in coords:
            continue
            
        rect = coords[field]
        if field == "Scope":
            print(f"[CHECK] IN LOOP: rect for Scope type={type(rect).__name__}, value={rect}")
            if isinstance(rect, dict):
                print(f"[WARN] rect is dict; keys={list(rect.keys())}, scope_layout={scope_layout}")
                rect = rect["long"] if estimated_lines >= 24 else rect["short"]
        
        # Template-specific starting font size for Scope
        if field == "Scope" and template_type == "standard" and scope_layout == "short":
            start_size = 15  # Standard template short scope: max 15pt
        else:
            # ✅ ADDED: Defensive check for font_starts dictionary
            if font_starts is None:
                print(f"⚠️ [CERTIFICATE] Font_starts dictionary is None - using default for field '{field}'")
                start_size = 30
            else:
                start_size = font_starts.get(field, 30)  # Use existing logic for other cases
        
        font_size = start_size
        

        

        
        # Reduce font size if it doesn't fit, but ensure minimum size
        while font_size >= 12:  # Increased minimum from 10 to 12
            text_height = get_text_height(text, font_size, fontname, rect.width)
            limit = rect.height if field != "Company Name" else rect.height * 2

            if text_height <= limit:
                break
            font_size -= 1
        

        
        if field == "Scope":
            # ✅ ADDED: Extract Excel adjustment and font size values for Scope
            scope_adjustment = parse_excel_adjustment(values.get("Scope Adjustment", ""))
            scope_font_size_adjustment = parse_excel_font_size(values.get("Scope Font Size", ""))
            
            # PowerPoint-style centering with automatic font size reduction
            original_font_size = font_size
            
            # ✅ ENHANCED: Use shared font calculation with binary search for line breaks
            
            def calculate_standard_font_size(text, rect, fontname, template_type, min_font_size):
                """Standard font calculation for text without line breaks"""
                font_size = original_font_size
                lines = []
                
                while font_size >= min_font_size:
                    # Process text with word wrapping
                    words = text.split()
                    current_line = ""
                    
                    for word in words:
                        is_bullet_point = any(word.startswith(indicator) for indicator in ['•', '>', '→', '▪', '▫', '*'])
                        
                        if is_bullet_point and current_line:
                            lines.append(current_line)
                            current_line = word
                            continue
                        
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        
                        if font_obj.text_length(test_line, font_size) <= rect.width:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        lines.append(current_line)
                    
                    # Calculate total height
                    if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                        line_height = font_size * 1.1
                    else:
                        line_height = font_size * 1.2
                    
                    total_height = len(lines) * line_height
                    
                    if total_height <= rect.height:
                        break
                    
                    font_size -= 1
                    lines = []  # Reset for next iteration
                
                return font_size, lines
            
        # ✅ USER CONTROL: Check if user wants to force font size (bypass optimization)
        force_font_size = values.get("Force Font Size", "").strip().lower()
        if force_font_size in ["true", "1", "yes", "force"]:
            print(f"🔍 [CERTIFICATE DEBUG] User requested force font size: {font_size}pt (bypassing optimization)")
            # Use the exact font size without optimization - just split text into lines
            lines = text.split('\n') if '\n' in text or '\r\n' in text else [text]
        else:
            # Use optimized font calculation
            min_font_size = 4  # Allow font size to go below 8pt if needed
            original_font_size = font_size
            font_size, lines = calculate_optimal_font_size_with_line_breaks(text, rect, fontname, template_type, min_font_size, font_size)
            print(f"🔍 [CERTIFICATE DEBUG] Binary search result: {font_size}pt (optimized from {original_font_size}pt)")
            
            # ✅ USER CONTROL: Apply Force Font Size as relative adjustment
            try:
                force_adjustment = float(force_font_size) if force_font_size else 0
                if force_adjustment != 0:
                    original_optimized_font = font_size
                    font_size += force_adjustment
                    print(f"🔍 [CERTIFICATE DEBUG] Force Font Size adjustment: {original_optimized_font}pt + {force_adjustment}pt = {font_size}pt")
            except (ValueError, TypeError):
                print(f"🔍 [CERTIFICATE DEBUG] Invalid Force Font Size value: '{force_font_size}' - treating as 0")
        
        # ✅ UPDATED: Apply Excel font size adjustment AFTER optimization (user override)
        if scope_font_size_adjustment != 0:
            optimized_scope_font = font_size
            font_size += scope_font_size_adjustment
            print(f"🔍 [CERTIFICATE DEBUG] Scope Font Size adjustment AFTER optimization: {optimized_scope_font}pt + {scope_font_size_adjustment}pt = {font_size}pt")
        
        # ✅ HORIZONTAL OVERFLOW FIX: Re-wrap lines if font size increased
        if scope_font_size_adjustment != 0:
            print(f"🔍 [SCOPE HORIZONTAL] Checking horizontal overflow at increased font size: {font_size}pt")
            
            font_obj = fitz.Font(fontname=fontname)
            rewrapped_lines = []
            rewrap_count = 0
            
            for line_idx, line in enumerate(lines):
                if not line.strip():
                    # Empty line - preserve as is
                    rewrapped_lines.append(line)
                    continue
                
                # Check if line fits at new font size
                line_width = font_obj.text_length(line, font_size)
                
                if line_width <= rect.width:
                    # Line still fits - keep as is
                    rewrapped_lines.append(line)
                else:
                    # Line overflows - re-wrap it
                    rewrap_count += 1
                    print(f"🔍 [SCOPE HORIZONTAL] Re-wrapping line {line_idx + 1}: '{line[:40]}...' (width: {line_width:.1f}pt > {rect.width:.1f}pt)")
                    
                    words = line.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        test_width = font_obj.text_length(test_line, font_size)
                        
                        if test_width <= rect.width:
                            current_line = test_line
                        else:
                            if current_line:
                                rewrapped_lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        rewrapped_lines.append(current_line)
            
            # Update lines with re-wrapped content
            original_line_count = len(lines)
            lines = rewrapped_lines
            new_line_count = len(lines)
            
            print(f"✅ [SCOPE HORIZONTAL] Re-wrapping complete: {original_line_count} → {new_line_count} lines ({rewrap_count} lines re-wrapped)")
        
        # Calculate final total height for overflow checking
        if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
            line_height = font_size * 1.1
        else:
            line_height = font_size * 1.2
        total_height = len(lines) * line_height

        # ✅ NEW: Binary search in font_utils handles overflow optimization automatically
        # No need for manual overflow handling - the shared function finds optimal font size

        # ✅ ENHANCED: Use optimized lines from font calculation
        # Replace all asterisks with bullet points for display in the optimized lines
        optimized_lines = []
        for line in lines:
            if line:  # Non-empty line
                display_line = line.replace('*', '•')
                if line != display_line:
                    print(f"🔄 [CERTIFICATE BULLET] Replaced '{line}' with '{display_line}'")
                optimized_lines.append(display_line)
            else:
                optimized_lines.append(line)  # Preserve empty lines
        
        lines = optimized_lines
        


        
        # Calculate total height and position vertically based on template type
        # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
        if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
            line_height = font_size * 1.1  # Tight spacing for large/logo templates
        else:  # standard templates
            line_height = font_size * 1.2  # Loose spacing for standard templates
        total_height = len(lines) * line_height
        
        
        # 🔍 DEBUG: Check for line breaks and template type
        has_line_breaks = '\n' in text or '\r\n' in text
        print(f"🔍 [CERTIFICATE DEBUG] Template type: {template_type}")
        print(f"🔍 [CERTIFICATE DEBUG] Has line breaks: {has_line_breaks}")
        print(f"🔍 [CERTIFICATE DEBUG] Rect coordinates: y0={rect.y0}, y1={rect.y1}, height={rect.height}")
        print(f"🔍 [CERTIFICATE DEBUG] Text height from binary search: {total_height}pt")
        print(f"🔍 [CERTIFICATE DEBUG] Font size from binary search: {font_size}pt")
        
        if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
            # Large/Logo template: start from top with no margin
            # If explicit line breaks, hard-code 7pt top offset (same as softcopy)
            if has_line_breaks:
                start_y = rect.y0 + 7 + scope_adjustment  # 7pt offset + Excel adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Large/Logo template with line breaks: start_y = {rect.y0} + 7 + {scope_adjustment} = {start_y}")
            else:
                start_y = rect.y0 + scope_adjustment  # Start at exact top of box + Excel adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Large/Logo template without line breaks: start_y = {rect.y0} + {scope_adjustment} = {start_y}")
            
            # Check if text would overflow bottom
            if start_y + total_height > rect.y1:
                # If overflow, adjust to fit within bounds
                old_start_y = start_y
                start_y = rect.y1 - total_height - 2  # 2pt margin from bottom
                print(f"⚠️ [CERTIFICATE DEBUG] Overflow detected! Adjusted start_y from {old_start_y} to {start_y}")
            else:
                print(f"✅ [CERTIFICATE DEBUG] No overflow: start_y={start_y}, end_y={start_y + total_height}, rect.y1={rect.y1}")
        else:
            # Standard template: top-align when explicit line breaks; otherwise center
            if has_line_breaks:
                # Hard-code 7pt top offset for explicit line breaks (same as softcopy) + Excel adjustment
                start_y = rect.y0 + 7 + scope_adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Standard template with line breaks: start_y = {rect.y0} + 7 + {scope_adjustment} = {start_y}")
            else:
                start_y = rect.y0 + (rect.height - total_height) / 2 + line_height/2 + scope_adjustment  # Adjust for baseline + Excel adjustment
                print(f"🔍 [CERTIFICATE DEBUG] Standard template without line breaks: start_y = {rect.y0} + {(rect.height - total_height) / 2 + line_height/2 + scope_adjustment} = {start_y}")
        
        # ✅ BULLET ALIGNMENT: Helper function to detect bullet lines
        def is_bullet_line(line):
            """Check if line starts with bullet character."""
            if not line or not line.strip():
                return False
            first_word = line.strip().split()[0] if line.strip().split() else ""
            return any(first_word.startswith(indicator) for indicator in ['•', '>', '→', '▪', '▫', '*'])
        
        # ✅ BULLET ALIGNMENT: Calculate smart left coordinate based on longest bullet sentence
        bullet_char_width = 0
        longest_bullet_line = None
        longest_word_count = 0
        font_obj = fitz.Font(fontname=fontname)
        
        # Find the longest bullet line by word count
        for line in lines:
            if is_bullet_line(line):
                word_count = len(line.strip().split())
                if word_count > longest_word_count:
                    longest_word_count = word_count
                    longest_bullet_line = line
        
        if longest_bullet_line:
            # Calculate bullet character width for the longest line
            first_word = longest_bullet_line.strip().split()[0]
            bullet_char_width = font_obj.text_length(first_word + " ", font_size)
            
            # Calculate what the leftmost coordinate would be if we center the longest line
            longest_line_width = font_obj.text_length(longest_bullet_line, font_size)
            center_x = (rect.x0 + rect.x1) / 2
            centered_leftmost = center_x - longest_line_width / 2
            
            # Use the leftmost coordinate from centering the longest line
            bullet_left_coord = centered_leftmost
            
            print(f"🔍 [CERTIFICATE BULLET] Longest bullet line: '{longest_bullet_line[:50]}...' ({longest_word_count} words)")
            print(f"🔍 [CERTIFICATE BULLET] Longest line width: {longest_line_width:.1f}pt")
            print(f"🔍 [CERTIFICATE BULLET] Centered leftmost coordinate: {centered_leftmost:.1f}pt")
            print(f"🔍 [CERTIFICATE BULLET] All bullets will align to: {bullet_left_coord:.1f}pt")
        else:
            bullet_left_coord = rect.x0 + bullet_char_width
            print(f"🔍 [CERTIFICATE BULLET] No bullets found, using default left coordinate: {bullet_left_coord:.1f}pt")
        
        print(f"🔍 [SCOPE BULLETS] Total lines: {len(lines)}")
        print(f"🔍 [SCOPE BULLETS] Bullet lines detected: {sum(1 for l in lines if is_bullet_line(l))}")
        print(f"🔍 [SCOPE BULLETS] Non-bullet lines: {sum(1 for l in lines if not is_bullet_line(l) and l.strip())}")
        
        # ✅ SIMPLIFIED: Scope rendering with consistent centering
        current_y = start_y
        
        # Render each line with bullet-aware alignment
        for i, line in enumerate(lines):
            if not line.strip():  # Skip empty lines
                current_y += line_height
                continue
            
            # Check if this is the last non-empty line
            is_last_line = i == len(lines) - 1 or all(not lines[j].strip() for j in range(i + 1, len(lines)))
            
            # NEW: Check if this is a bullet line
            is_bullet = is_bullet_line(line)
            
            if is_bullet:
                # BULLET LINE: Left-align with smart left coordinate from longest line centering
                start_x = bullet_left_coord
                
                # Split bullet character from text
                first_word = line.strip().split()[0]  # Bullet character (•, >, →, etc.)
                rest_of_text = line[len(first_word):].strip()  # Text after bullet
                
                # Calculate bullet font size (30% larger than normal)
                bullet_font_size = font_size * 1.3
                
                # Render bullet character with larger font
                safe_insert_text(page, (start_x, current_y), first_word, 
                               fontsize=bullet_font_size, fontname=fontname, color=color)
                
                # Calculate position for text after bullet
                font_obj = fitz.Font(fontname=fontname)
                bullet_width = font_obj.text_length(first_word + " ", bullet_font_size)
                text_start_x = start_x + bullet_width
                
                # Render text with normal font size
                if rest_of_text:
                    safe_insert_text(page, (text_start_x, current_y), rest_of_text, 
                                   fontsize=font_size, fontname=fontname, color=color)
                
                print(f"🔍 [CERTIFICATE BULLET] Line {i+1} left-aligned to {start_x:.1f}pt: '{first_word}' ({bullet_font_size:.1f}pt) + '{rest_of_text[:30]}...' ({font_size}pt)")
            elif is_last_line:
                # ✅ LAST LINE: Center align for balanced appearance
                center_x = (rect.x0 + rect.x1) / 2
                
                # ✅ ENHANCED: Use mixed format text rendering for bold detection
                if '**' in line or '__' in line:
                    # Calculate total width for centering
                    segments = process_bold_text(line)
                    total_width = 0
                    for segment_text, _, _ in segments:
                        if segment_text:
                            font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                            total_width += font_obj.text_length(segment_text, font_size)
                    
                    start_x = center_x - total_width / 2
                    render_mixed_format_text(page, (start_x, current_y), line, font_size, color)
                else:
                    # Standard rendering for non-bold text
                    font_obj = fitz.Font(fontname=fontname)
                    text_width = font_obj.text_length(line, font_size)
                    start_x = center_x - text_width / 2
                    
                    safe_insert_text(
                        page,
                        (start_x, current_y),
                        line,
                        fontsize=font_size,
                        fontname=fontname,
                        color=color
                    )
                # ✅ Safe string handling for debug output
                safe_line = str(line) if line is not None else ""
                print(f"🔍 [CERTIFICATE] Last line centered: '{safe_line[:50]}{'...' if len(safe_line) > 50 else ''}'")
            else:
                # ✅ INTERMEDIATE LINES: Center align for consistency
                center_x = (rect.x0 + rect.x1) / 2
                
                # ✅ ENHANCED: Use mixed format text rendering for bold detection
                if '**' in line or '__' in line:
                    # Calculate total width for centering
                    segments = process_bold_text(line)
                    total_width = 0
                    for segment_text, _, _ in segments:
                        if segment_text:
                            font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                            total_width += font_obj.text_length(segment_text, font_size)
                    
                    start_x = center_x - total_width / 2
                    render_mixed_format_text(page, (start_x, current_y), line, font_size, color)
                else:
                    # Standard rendering for non-bold text
                    font_obj = fitz.Font(fontname=fontname)
                    text_width = font_obj.text_length(line, font_size)
                    start_x = center_x - text_width / 2
                    
                    safe_insert_text(
                        page,
                        (start_x, current_y),
                        line,
                        fontsize=font_size,
                        fontname=fontname,
                        color=color
                    )
                
                # ✅ Safe string handling for debug output
                safe_line = str(line) if line is not None else ""
                print(f"🔍 [CERTIFICATE] Line {i+1} centered: '{safe_line[:50]}{'...' if len(safe_line) > 50 else ''}'")
            
            # Update current_y consistently for all lines
            # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
            if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_nonaccredited_other", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                current_y += font_size * 1.1  # Tight spacing for large/logo templates
            else:  # standard templates
                current_y += font_size * 1.2  # Loose spacing for standard templates
        
        print(f"🎯 [SCOPE SUMMARY] Final rendering method: CENTERED")
        print(f"🎯 [SCOPE SUMMARY] Scope field '{field}' processed successfully")
        print(f"🔍 [CERTIFICATE DEBUG] ===== SCOPE ANALYSIS COMPLETE =====")
        print(f"🔍 [CERTIFICATE DEBUG] Binary search result: font_size={font_size:.1f}pt, lines={len(lines)}")
        print(f"🔍 [CERTIFICATE DEBUG] Line height: {line_height:.1f}pt")
        print(f"🔍 [CERTIFICATE DEBUG] Total text height: {total_height:.1f}pt")
        print(f"🔍 [CERTIFICATE DEBUG] Available rect height: {rect.height:.1f}pt")
        print(f"🔍 [CERTIFICATE DEBUG] Rect boundaries: y0={rect.y0:.1f}, y1={rect.y1:.1f}")
        print(f"🔍 [CERTIFICATE DEBUG] Final coordinates used: start_y={start_y:.1f}, end_y={start_y + total_height:.1f}")
        print(f"🔍 [CERTIFICATE DEBUG] Space utilization: {((start_y + total_height - rect.y0) / rect.height) * 100:.1f}%")
        print(f"🔍 [CERTIFICATE DEBUG] Remaining space: {rect.y1 - (start_y + total_height):.1f}pt")
        print(f"🔍 [CERTIFICATE DEBUG] Binary search working: {'YES' if font_size < 20 else 'NO'} (font reduced from 20pt to {font_size:.1f}pt)")
        print(f"🔍 [CERTIFICATE DEBUG] ===== END SCOPE ANALYSIS =====")

    # ✅ ADDED: Insert logo if available and using any logo template type
    if logo_image and template_type.startswith("logo"):
        try:
            # ✅ ADDED: Defensive check for logo_coords
            if logo_coords is None:
                print(f"⚠️ [CERTIFICATE] Logo_coords is None - skipping logo insertion")
            else:
                # Get logo coordinates from logo_coords
                logo_rect = logo_coords.get("logo")
                if logo_rect:
                    # Use the smart positioning function directly with the loaded logo_image
                    insert_logo_with_smart_positioning(page, logo_image, logo_rect)
                    print(f"✅ [CERTIFICATE] Logo inserted for template type: {template_type}")
                else:
                    print(f"⚠️ [CERTIFICATE] Logo coordinates not found in logo_coords for template: {template_type}")
        except Exception as logo_insert_error:
            print(f"❌ [CERTIFICATE] Error inserting logo: {logo_insert_error}")

    # ✅ ADDED: Process Extra Line field
    extra_line_text = values.get("Extra Line", "").strip()
    if extra_line_text:
        print(f"🔍 [CERTIFICATE] Processing Extra Line: '{extra_line_text}'")
        
        # Calculate Extra Line position (0pt gap below scope)
        # Use the same scope_rect that was used for scope rendering
        if template_type in ["large", "large_eco", "large_nonaccredited", "large_other", "large_other_eco", "large_other_nonaccredited", "large_nonaccredited_other"]:
            scope_rect = coords["Scope"]  # Single rectangle for large templates
        else:
            # For standard/logo templates, use the stored original coordinates
            if estimated_lines >= 24:
                scope_rect = original_scope_coords["long"]
            else:
                scope_rect = original_scope_coords["short"]
        
        # Calculate Extra Line position based on content length
        scope_text = values.get("Scope", "")
        scope_words = len(scope_text.split())
        estimated_lines = max(1, (scope_words * 8) // 60)
        
        if estimated_lines < 24:
            extra_line_y = scope_rect.y1 + 25  # 25pt gap below scope for <24 lines
        else:
            extra_line_y = scope_rect.y1  # 0pt gap - directly below scope for ≥24 lines
        
        # Create Extra Line rectangle
        extra_line_rect = fitz.Rect(
            scope_rect.x0,      # Same x0 as scope
            extra_line_y,       # Dynamic gap below scope
            scope_rect.x1,      # Same x1 as scope  
            extra_line_y + 10   # 10pt height for text
        )
        
        # Render Extra Line text with center alignment and bold font
        try:
            if '**' in extra_line_text or '__' in extra_line_text:
                # Use mixed format rendering for bold text with center alignment
                render_mixed_format_text(page, (extra_line_rect.x0, extra_line_rect.y0), extra_line_text, 12, (0, 0, 0), extra_line_rect.width)
            else:
                # Center-aligned bold text rendering
                center_x = (extra_line_rect.x0 + extra_line_rect.x1) / 2
                font_obj = fitz.Font(fontname="Times-Bold")
                text_width = font_obj.text_length(extra_line_text, 12)
                start_x = center_x - text_width / 2
                
                safe_insert_text(
                    page,
                    (start_x, extra_line_rect.y0),
                    extra_line_text,
                    fontsize=12,
                    fontname="Times-Bold",
                    color=(0, 0, 0)
                )
            
            print(f"🔍 [CERTIFICATE] Extra Line rendered at: {extra_line_rect}")
        except Exception as extra_line_error:
            print(f"❌ [CERTIFICATE] Error rendering Extra Line: {extra_line_error}")
            print(f"🔍 [CERTIFICATE] Extra Line coordinates: {extra_line_rect}")
            print(f"🔍 [CERTIFICATE] Extra Line text: '{extra_line_text}'")
            # Continue without Extra Line rather than failing completely
    else:
        print(f"🔍 [CERTIFICATE] No Extra Line - skipping")

    # ✅ ADDED: Robust return structure - always save and return
    try:
        doc.save(output_pdf_path)
        doc.close()
        
        print(f"[CERTIFICATE] Certificate PDF generated successfully: {output_pdf_path}")
        
        # Return tracking information
        return {
            "success": True,
            "output_path": output_pdf_path,
            "overflow_warnings": overflow_warnings,
            "template_type": template_type
        }
    except Exception as save_error:
        print(f"❌ [CERTIFICATE] Error saving PDF: {save_error}")
        # Still return a result dict even if save fails
        return {
            "success": False,
            "error": f"Failed to save PDF: {save_error}",
            "output_path": output_pdf_path,
            "overflow_warnings": overflow_warnings,
            "template_type": template_type
        }



